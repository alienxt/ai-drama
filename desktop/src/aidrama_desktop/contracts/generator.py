from __future__ import annotations

import hashlib
import json
import os
import re
import secrets
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import date, timedelta
from decimal import Decimal, InvalidOperation
from pathlib import Path
from typing import Callable
from xml.etree import ElementTree

from docx import Document


CONTRACT_MEDIA_PLATFORMS = ("WECHAT_VIDEO",)
CONTRACT_TEMPLATE_TYPES = ("cost", "purchase", "rights")
CONTRACT_TEMPLATE_TYPE_LABELS = {
    "cost": "成本合同",
    "purchase": "购买合同",
    "rights": "权利声明",
}
CONTRACT_PARTY_FIELD_LABELS = {
    "buyer": "买方/甲方",
    "seller": "卖方/乙方",
}
CONTRACT_PLATFORM_TEMPLATE_TYPES = {
    "WECHAT_VIDEO": ("cost", "purchase", "rights"),
    "TIKTOK": ("purchase",),
    "DOUYIN": ("purchase",),
}
CONTRACT_MATERIAL_CACHE_VERSION = 4
DOUBLE_DATE_FOOTER_SPACER = " " * 20
DOUBLE_DATE_FOOTER_RE = re.compile(
    r"(日期\s*[:：]\s*)\{\{\s*date\s*\}\}(\s+)(日期\s*[:：]\s*)\{\{\s*date\s*\}\}"
)
PURCHASE_STAMP_VERTICAL_SHIFT_EMU = 0
PURCHASE_SELLER_STAMP_VERTICAL_SHIFT_EMU = 0
DOCX_DOCUMENT_XML = "word/document.xml"
DOCX_WORD_XML_PREFIX = "word/"
DOCX_WORD_XML_SUFFIX = ".xml"
WORDML_NAMESPACES = {
    "wpc": "http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "o": "urn:schemas-microsoft-com:office:office",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "v": "urn:schemas-microsoft-com:vml",
    "wp14": "http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "w10": "urn:schemas-microsoft-com:office:word",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
    "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
    "wpi": "http://schemas.microsoft.com/office/word/2010/wordprocessingInk",
    "wne": "http://schemas.microsoft.com/office/word/2006/wordml",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "wpsCustomData": "http://www.wps.cn/officeDocument/2013/wpsCustomData",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}

for prefix, namespace in WORDML_NAMESPACES.items():
    ElementTree.register_namespace(prefix, namespace)


def contract_template_key(platform: str, contract_type: str) -> str:
    return f"{platform.lower()}:{contract_type}"


def contract_party_key(platform: str, party: str) -> str:
    return f"{platform.lower()}:{party}"


def required_contract_template_types(platform: str) -> tuple[tuple[str, str], ...]:
    types = CONTRACT_PLATFORM_TEMPLATE_TYPES.get(platform, ("purchase",))
    return tuple((contract_type, CONTRACT_TEMPLATE_TYPE_LABELS[contract_type]) for contract_type in types)


def required_contract_party_fields(platform: str) -> tuple[tuple[str, str], ...]:
    if platform != "WECHAT_VIDEO":
        return ()
    return tuple(CONTRACT_PARTY_FIELD_LABELS.items())


def all_required_contract_templates_configured(templates: dict[str, Path | str | None], platform: str) -> bool:
    for contract_type, _label in required_contract_template_types(platform):
        value = templates.get(contract_template_key(platform, contract_type))
        if not value or not Path(value).exists():
            return False
    return True


@dataclass(frozen=True)
class ContractRenderInput:
    contract_type: str
    drama_title: str
    episode_count: str
    price: str
    buyer: str
    seller: str
    sign_date: str = ""
    start_date: str = ""
    episode_minutes: str = ""
    agreement_number: str = ""

    def placeholders(self) -> dict[str, str]:
        start_date = self.start_date or generate_contract_start_date(
            self.sign_date,
            f"{self.contract_type}:{self.drama_title}:{self.agreement_number}",
        )
        return {
            "contractType": self.contract_type,
            "agreementNumber": self.agreement_number,
            "dramaTitle": self.drama_title,
            "episodeCount": self.episode_count,
            "episodeMinutes": self.episode_minutes,
            "price": self.price,
            "halfPrice": format_half_price(self.price),
            "buyer": self.buyer,
            "seller": self.seller,
            "date": format_contract_date_no_wrap(self.sign_date),
            "startDate": format_contract_date_no_wrap(start_date),
        }


@dataclass(frozen=True)
class ContractMaterial:
    contract_type: str
    label: str
    docx_path: Path
    image_paths: list[Path]


@dataclass(frozen=True)
class ContractMaterialBundle:
    materials: list[ContractMaterial]

    def metadata(self) -> dict[str, object]:
        result: dict[str, object] = {"contractMaterials": self.materials}
        purchase_images: list[Path] = []
        rights_images: list[Path] = []
        for material in self.materials:
            if material.contract_type == "purchase":
                result["purchaseContractDocx"] = material.docx_path
                result["purchaseContractImages"] = material.image_paths
                purchase_images = material.image_paths
            elif material.contract_type == "cost":
                result["costContractDocx"] = material.docx_path
                result["costContractImages"] = material.image_paths
                result["costConfigReportImages"] = material.image_paths
            elif material.contract_type == "rights":
                result["rightsStatementDocx"] = material.docx_path
                result["rightsStatementImages"] = material.image_paths
                rights_images = material.image_paths
        if purchase_images or rights_images:
            result["buyDramaContractImages"] = [*purchase_images, *rights_images]
        return result


def default_contract_templates() -> dict[str, Path | None]:
    return {
        contract_template_key(platform, contract_type): None
        for platform in CONTRACT_MEDIA_PLATFORMS
        for contract_type in CONTRACT_TEMPLATE_TYPES
    }


def default_contract_config() -> dict[str, Path | str | None]:
    return {
        **default_contract_templates(),
        **{
            contract_party_key(platform, party): ""
            for platform in CONTRACT_MEDIA_PLATFORMS
            for party in CONTRACT_PARTY_FIELD_LABELS
        },
    }


class ContractConfigStore:
    def __init__(self, path: Path):
        self.path = path

    def load(self) -> dict[str, Path | str | None]:
        if not self.path.exists():
            return default_contract_config()
        try:
            data = json.loads(self.path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            return default_contract_config()
        templates = default_contract_config()
        for key in default_contract_templates():
            value = data.get(key)
            if isinstance(value, str) and value.strip():
                templates[key] = Path(value)
        for platform in CONTRACT_MEDIA_PLATFORMS:
            for party in CONTRACT_PARTY_FIELD_LABELS:
                key = contract_party_key(platform, party)
                value = data.get(key)
                if isinstance(value, str):
                    templates[key] = value.strip()
        for legacy_key in CONTRACT_TEMPLATE_TYPES:
            value = data.get(legacy_key)
            migrated_key = contract_template_key("WECHAT_VIDEO", legacy_key)
            if templates.get(migrated_key) is None and isinstance(value, str) and value.strip():
                templates[migrated_key] = Path(value)
        return templates

    def save(self, templates: dict[str, Path | str | None]) -> None:
        self.path.parent.mkdir(parents=True, exist_ok=True)
        data: dict[str, str] = {}
        for key in default_contract_config():
            value = templates.get(key)
            data[key] = str(value) if value else ""
        self.path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def safe_contract_filename(filename: str) -> str:
    return re.sub(r"[^0-9A-Za-z\u4e00-\u9fff._-]+", "_", filename).strip("_") or "contract"


def copy_contract_template(source: Path, template_dir: Path, name: str) -> Path:
    if source.suffix.lower() != ".docx":
        raise ValueError("请选择 .docx 格式的 Word 模板。")
    if not source.exists():
        raise FileNotFoundError(f"模板不存在：{source}")
    template_dir.mkdir(parents=True, exist_ok=True)
    target = template_dir / f"{safe_contract_filename(name)}.docx"
    shutil.copy2(source, target)
    return target


def build_contract_template_download_path(target_dir: Path, key: str, template: dict[str, object]) -> Path:
    source_name = str(template.get("name") or template.get("fileName") or key)
    file_name = safe_contract_filename(source_name)
    if not file_name.lower().endswith(".docx"):
        file_name = f"{file_name}.docx"
    prefix = safe_contract_filename(key)
    return target_dir / f"{prefix}-{file_name}"


def replace_contract_text(text: str, values: dict[str, str]) -> str:
    def replace(match: re.Match[str]) -> str:
        key = match.group(1).strip()
        return values.get(key, "")

    return re.sub(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}", replace, text)


def format_contract_date_compact(value: str = "") -> str:
    return format_contract_date(value).replace(" ", "")


def format_half_price(value: str = "") -> str:
    try:
        amount = Decimal(str(value).strip() or "0") / Decimal("2")
    except (InvalidOperation, ValueError):
        return ""
    normalized = amount.normalize()
    if normalized == normalized.to_integral():
        return str(normalized.quantize(Decimal("1")))
    return format(normalized, "f").rstrip("0").rstrip(".")


def generate_agreement_number(value: str = "") -> str:
    parsed = parse_contract_date(value) or date.today()
    return f"HZ-{parsed.year:04d}-{parsed.month:02d}-{secrets.randbelow(1_000_000):06d}"


def generate_contract_start_date(value: str = "", seed: str | None = None) -> str:
    parsed = parse_contract_date(value) or date.today()
    if seed is None:
        offset_days = 30 + secrets.randbelow(11)
    else:
        digest = hashlib.sha256(seed.encode("utf-8")).digest()
        offset_days = 30 + digest[0] % 11
    return (parsed - timedelta(days=offset_days)).isoformat()


def format_contract_date_short(value: str = "") -> str:
    compact = format_contract_date_compact(value)
    match = re.fullmatch(r"(\d{4})年(\d{2})月(\d{2})日", compact)
    if not match:
        return compact
    return ".".join(match.groups())


def format_contract_date_no_wrap(value: str = "") -> str:
    return format_contract_date_compact(value)


def format_contract_date(value: str = "") -> str:
    parsed = parse_contract_date(value)
    if parsed is None:
        return value.strip()
    return f"{parsed.year:04d} 年 {parsed.month:02d} 月 {parsed.day:02d} 日"


def parse_contract_date(value: str = "") -> date | None:
    clean = value.strip()
    if not clean:
        return date.today()
    normalized = (
        clean.replace("年", "-")
        .replace("月", "-")
        .replace("日", "")
        .replace("/", "-")
        .replace(".", "-")
    )
    normalized = re.sub(r"\s+", "", normalized)
    match = re.fullmatch(r"(\d{4})-(\d{1,2})-(\d{1,2})", normalized)
    if not match:
        return None
    year, month, day = (int(part) for part in match.groups())
    try:
        return date(year, month, day)
    except ValueError:
        return None


def replace_double_date_footer_text(text: str, date_value: str) -> str:
    def replace(match: re.Match[str]) -> str:
        return (
            f"{match.group(1)}{date_value}"
            f"{DOUBLE_DATE_FOOTER_SPACER}"
            f"{match.group(3)}{date_value}"
        )

    return DOUBLE_DATE_FOOTER_RE.sub(replace, text)


def replace_paragraph_text(paragraph, values: dict[str, str]) -> None:
    original = paragraph.text
    prepared = replace_double_date_footer_text(original, values.get("date", ""))
    replaced = replace_contract_text(prepared, values)
    if replaced == original:
        return
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = replaced
    else:
        paragraph.add_run(replaced)


def render_contract_docx(
    template: Path,
    output: Path,
    data: ContractRenderInput,
    normalize_for_rendering: bool = True,
) -> Path:
    if not template.exists():
        raise FileNotFoundError(f"请先配置 Word 模板：{template}")
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document(template)
    values = data.placeholders()
    for paragraph in document.paragraphs:
        replace_paragraph_text(paragraph, values)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph_text(paragraph, values)
    document.save(output)
    if normalize_for_rendering:
        normalize_contract_docx_for_rendering(output)
    return output


def build_contract_output_path(output_dir: Path, data: ContractRenderInput) -> Path:
    filename = safe_contract_filename(f"{data.contract_type}-{data.drama_title}-{data.sign_date or date.today().isoformat()}")
    return output_dir / f"{filename}.docx"


def render_contract_material_bundle(
    templates: dict[str, Path | str | None],
    platform: str,
    output_dir: Path,
    data_factory: Callable[[str, str], ContractRenderInput],
    image_converter: Callable[[Path, Path, str | None], list[Path]] | None = None,
    soffice_path: str | None = None,
) -> ContractMaterialBundle:
    if not all_required_contract_templates_configured(templates, platform):
        missing = [
            label
            for contract_type, label in required_contract_template_types(platform)
            if not templates.get(contract_template_key(platform, contract_type))
            or not Path(templates[contract_template_key(platform, contract_type)]).exists()
        ]
        raise RuntimeError(f"请先在合同配置中配置：{'、'.join(missing)}。")

    materials: list[ContractMaterial] = []
    for contract_type, label in required_contract_template_types(platform):
        template = Path(templates[contract_template_key(platform, contract_type)])
        data = data_factory(contract_type, label)
        normalize_for_rendering = contract_type == "cost"
        docx_path = build_contract_output_path(output_dir / "docx", data)
        image_stem = safe_contract_filename(f"{contract_type}-{data.drama_title}")
        cache_key = build_contract_material_cache_key(
            template,
            platform,
            contract_type,
            label,
            data,
            normalize_for_rendering,
        )
        manifest_path = build_contract_material_cache_path(output_dir / "cache", image_stem)
        cached_material = load_cached_contract_material(manifest_path, cache_key, contract_type, label)
        if cached_material:
            materials.append(cached_material)
            continue

        docx_path = render_contract_docx(
            template,
            docx_path,
            data,
            normalize_for_rendering=normalize_for_rendering,
        )
        image_paths = convert_contract_docx_images(
            contract_type,
            docx_path,
            output_dir / "images",
            image_stem,
            image_converter,
            soffice_path=soffice_path,
        )
        if contract_type == "purchase" and len(image_paths) > 1:
            image_paths = [
                merge_pngs_vertically(
                    image_paths,
                    output_dir / "images" / f"{safe_contract_filename(image_stem)}.png",
                )
            ]
        if not image_paths or any(not path.exists() for path in image_paths):
            raise RuntimeError(f"{label}图片生成失败。")
        material = ContractMaterial(contract_type, label, docx_path, image_paths)
        write_contract_material_cache(manifest_path, cache_key, material)
        materials.append(material)
    return ContractMaterialBundle(materials)


def convert_contract_docx_images(
    contract_type: str,
    docx_path: Path,
    image_dir: Path,
    image_stem: str,
    image_converter: Callable[[Path, Path, str | None], list[Path]] | None = None,
    soffice_path: str | None = None,
) -> list[Path]:
    if image_converter is not None:
        return image_converter(docx_path, image_dir, image_stem)
    if contract_type == "rights":
        quicklook = quicklook_docx_to_image(docx_path, image_dir, image_stem)
        if quicklook:
            return [quicklook]
    allow_single_page_fallback = contract_type != "purchase"
    return export_contract_docx_images(
        docx_path,
        image_dir,
        image_stem,
        soffice_path=soffice_path,
        allow_quicklook_fallback=allow_single_page_fallback,
        allow_single_page_pdf_fallback=allow_single_page_fallback,
    )


def build_contract_material_cache_path(cache_dir: Path, image_stem: str) -> Path:
    return cache_dir / f"{safe_contract_filename(image_stem)}.json"


def build_contract_material_cache_key(
    template: Path,
    platform: str,
    contract_type: str,
    label: str,
    data: ContractRenderInput,
    normalize_for_rendering: bool,
) -> dict[str, object]:
    resolved_template = template.resolve()
    return {
        "version": CONTRACT_MATERIAL_CACHE_VERSION,
        "platform": platform,
        "contractType": contract_type,
        "label": label,
        "normalizeForRendering": normalize_for_rendering,
        "templatePath": str(resolved_template),
        "templateSha256": file_sha256(resolved_template),
        "placeholders": data.placeholders(),
        "signDate": data.sign_date,
    }


def load_cached_contract_material(
    manifest_path: Path,
    cache_key: dict[str, object],
    contract_type: str,
    label: str,
) -> ContractMaterial | None:
    try:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return None
    if manifest.get("cacheKey") != cache_key:
        return None
    docx_path = Path(str(manifest.get("docxPath") or ""))
    image_paths = [Path(str(path)) for path in manifest.get("imagePaths") or []]
    if not docx_path.exists() or not image_paths or any(not path.exists() for path in image_paths):
        return None
    return ContractMaterial(contract_type, label, docx_path, image_paths)


def write_contract_material_cache(
    manifest_path: Path,
    cache_key: dict[str, object],
    material: ContractMaterial,
) -> None:
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(
        json.dumps(
            {
                "cacheKey": cache_key,
                "docxPath": str(material.docx_path),
                "imagePaths": [str(path) for path in material.image_paths],
            },
            ensure_ascii=False,
            indent=2,
            sort_keys=True,
        ),
        encoding="utf-8",
    )


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for chunk in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def normalize_contract_docx_for_rendering(docx_path: Path) -> bool:
    try:
        with zipfile.ZipFile(docx_path) as docx:
            xml_members = [
                info.filename
                for info in docx.infolist()
                if info.filename.startswith(DOCX_WORD_XML_PREFIX)
                and info.filename.endswith(DOCX_WORD_XML_SUFFIX)
            ]
            member_contents = {member: docx.read(member) for member in xml_members}
    except zipfile.BadZipFile:
        return False

    changed_members: dict[str, bytes] = {}
    for member, content in member_contents.items():
        try:
            root = ElementTree.fromstring(content)
        except ElementTree.ParseError:
            continue
        changed = False
        changed |= _normalize_run_spacing(root)
        if changed:
            changed_members[member] = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)

    if not changed_members:
        return False

    _replace_docx_members(docx_path, changed_members)
    return True


def _normalize_run_spacing(root: ElementTree.Element) -> bool:
    changed = False
    for spacing in root.findall(".//w:rPr/w:spacing", WORDML_NAMESPACES):
        value = spacing.get(_word_attr("val"))
        if not value:
            continue
        try:
            numeric = int(value)
        except ValueError:
            continue
        if numeric < 0:
            spacing.set(_word_attr("val"), "0")
            changed = True
    return changed


def _normalize_cjk_fonts(root: ElementTree.Element) -> bool:
    changed = False
    for fonts in root.findall(".//w:rFonts", WORDML_NAMESPACES):
        for name in ("eastAsia", "ascii", "hAnsi", "cs"):
            attr = _word_attr(name)
            value = fonts.get(attr)
            if value and _is_unstable_render_font(value):
                fonts.set(attr, "Songti SC")
                changed = True
        for theme_name in ("eastAsiaTheme", "asciiTheme", "hAnsiTheme", "cstheme"):
            attr = _word_attr(theme_name)
            if attr in fonts.attrib:
                del fonts.attrib[attr]
                changed = True
    return changed


def _is_unstable_render_font(name: str) -> bool:
    normalized = name.lower().replace(" ", "")
    return any(
        marker in normalized
        for marker in (
            "仿宋",
            "宋体",
            "黑体",
            "楷体",
            "等线",
            "gb2312",
            "minor",
        )
    )


def normalize_purchase_contract_layout(docx_path: Path) -> bool:
    try:
        with zipfile.ZipFile(docx_path) as docx:
            document_xml = docx.read(DOCX_DOCUMENT_XML)
    except (KeyError, zipfile.BadZipFile):
        return False

    try:
        root = ElementTree.fromstring(document_xml)
    except ElementTree.ParseError:
        return False

    anchors = root.findall(".//wp:anchor", WORDML_NAMESPACES)
    if not anchors:
        return False

    stamp_anchors = [
        (anchor, description)
        for anchor in anchors
        if _is_purchase_stamp_description(description := _drawing_description(anchor))
    ]
    if not stamp_anchors:
        return False

    stamp_anchor_ids = {id(anchor) for anchor, _description in stamp_anchors}
    non_stamp_max_relative_height = max(
        (_int_attr(anchor, "relativeHeight") for anchor in anchors if id(anchor) not in stamp_anchor_ids),
        default=0,
    )
    next_stamp_height = non_stamp_max_relative_height + 4096
    changed = False

    for anchor, description in stamp_anchors:
        changed |= _set_anchor_attr(anchor, "behindDoc", "0")
        changed |= _set_anchor_attr(anchor, "allowOverlap", "1")
        changed |= _set_anchor_attr(anchor, "relativeHeight", str(next_stamp_height))
        next_stamp_height += 4096
        changed |= _ensure_anchor_shifted_up(anchor, _purchase_stamp_vertical_shift_emu(description))

    if not changed:
        return False

    updated_xml = ElementTree.tostring(root, encoding="utf-8", xml_declaration=True)
    _replace_docx_member(docx_path, DOCX_DOCUMENT_XML, updated_xml)
    return True


def _drawing_description(anchor: ElementTree.Element) -> str:
    descriptions: list[str] = []
    for tag in ("wp:docPr", ".//pic:cNvPr"):
        element = anchor.find(tag, WORDML_NAMESPACES)
        if element is None:
            continue
        descriptions.extend(
            value
            for value in (element.get("descr"), element.get("name"))
            if value
        )
    return " ".join(descriptions)


def _is_purchase_stamp_description(description: str) -> bool:
    return "盖章" in description or "公章" in description or ("乙方" in description and "签名" not in description)


def _purchase_stamp_vertical_shift_emu(description: str) -> int:
    if "乙方" in description:
        return PURCHASE_SELLER_STAMP_VERTICAL_SHIFT_EMU
    return PURCHASE_STAMP_VERTICAL_SHIFT_EMU


def _set_anchor_attr(anchor: ElementTree.Element, name: str, value: str) -> bool:
    if anchor.get(name) == value:
        return False
    anchor.set(name, value)
    return True


def _ensure_anchor_shifted_up(anchor: ElementTree.Element, shift_emu: int) -> bool:
    if shift_emu <= 0:
        return False
    position = anchor.find("wp:positionV/wp:posOffset", WORDML_NAMESPACES)
    if position is None or position.text is None:
        return False
    try:
        current = int(position.text)
    except ValueError:
        return False
    if current < 0:
        return False
    target = current - shift_emu
    if target == current:
        return False
    position.text = str(target)
    return True


def _int_attr(element: ElementTree.Element, name: str) -> int:
    try:
        return int(element.get(name) or "0")
    except ValueError:
        return 0


def _replace_docx_member(docx_path: Path, member_name: str, content: bytes) -> None:
    _replace_docx_members(docx_path, {member_name: content})


def _replace_docx_members(docx_path: Path, members: dict[str, bytes]) -> None:
    with tempfile.NamedTemporaryFile(dir=docx_path.parent, suffix=".docx", delete=False) as temp_file:
        temp_path = Path(temp_file.name)
    try:
        with zipfile.ZipFile(docx_path) as source, zipfile.ZipFile(temp_path, "w") as target:
            for info in source.infolist():
                data = members.get(info.filename, source.read(info.filename))
                replacement = zipfile.ZipInfo(info.filename, info.date_time)
                replacement.comment = info.comment
                replacement.extra = info.extra
                replacement.internal_attr = info.internal_attr
                replacement.external_attr = info.external_attr
                replacement.compress_type = zipfile.ZIP_DEFLATED
                target.writestr(replacement, data)
        temp_path.replace(docx_path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def _word_attr(name: str) -> str:
    return f"{{{WORDML_NAMESPACES['w']}}}{name}"


def export_contract_docx_images(
    docx_path: Path,
    image_dir: Path,
    image_stem: str | None = None,
    *,
    soffice_path: str | None = None,
    allow_quicklook_fallback: bool = True,
    allow_single_page_pdf_fallback: bool = True,
) -> list[Path]:
    image_dir.mkdir(parents=True, exist_ok=True)
    stem = image_stem or docx_path.stem
    try:
        pdf_path = convert_docx_to_pdf(docx_path, image_dir, soffice_path=soffice_path)
    except RuntimeError:
        if not allow_quicklook_fallback:
            raise
        quicklook = quicklook_docx_to_image(docx_path, image_dir, stem)
        if quicklook:
            return [quicklook]
        raise
    try:
        return convert_pdf_to_pngs(
            pdf_path,
            image_dir,
            stem,
            allow_single_page_fallback=allow_single_page_pdf_fallback,
        )
    except RuntimeError:
        if not allow_quicklook_fallback:
            raise
        quicklook = quicklook_docx_to_image(docx_path, image_dir, stem)
        if quicklook:
            return [quicklook]
        raise


def convert_docx_to_pdf(docx_path: Path, output_dir: Path, *, soffice_path: str | None = None) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    soffice_options = find_command_options(
        "soffice",
        [
            soffice_path or "",
            os.environ.get("AIDRAMA_SOFFICE_PATH", ""),
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "/opt/homebrew/bin/soffice",
            "/usr/local/bin/soffice",
            "C:/Program Files/LibreOffice/program/soffice.exe",
            str(Path.home() / ".cache/codex-runtimes/codex-primary-runtime/dependencies/bin/soffice"),
        ],
        prefer_candidates=bool(str(soffice_path or "").strip() and str(soffice_path).strip() != "soffice"),
    )
    if not soffice_options:
        raise RuntimeError("无法将 Word 合同转换为图片，请先安装 LibreOffice。")
    errors: list[str] = []
    for soffice in soffice_options:
        with tempfile.TemporaryDirectory(prefix="aidrama-libreoffice-") as profile_dir:
            command = [
                soffice,
                f"-env:UserInstallation={Path(profile_dir).as_uri()}",
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "pdf:writer_pdf_Export",
                "--outdir",
                str(output_dir),
                str(docx_path),
            ]
            try:
                run_command(command, f"Word 合同转 PDF 失败（LibreOffice：{soffice}）")
            except RuntimeError as exception:
                errors.append(str(exception))
                continue
        pdf_path = output_dir / f"{docx_path.stem}.pdf"
        if not pdf_path.exists():
            matches = list(output_dir.glob(f"{docx_path.stem}*.pdf"))
            if matches:
                pdf_path = matches[0]
        if pdf_path.exists():
            return pdf_path
        errors.append(f"Word 合同转 PDF 后没有找到输出文件（LibreOffice：{soffice}）。")
    detail = "\n".join(errors[-3:])
    raise RuntimeError(f"无法使用 LibreOffice 转换 Word 合同：\n{detail}" if detail else "无法使用 LibreOffice 转换 Word 合同。")


def convert_pdf_to_pngs(
    pdf_path: Path,
    image_dir: Path,
    image_stem: str,
    *,
    allow_single_page_fallback: bool = True,
) -> list[Path]:
    pdftoppm = find_command(
        "pdftoppm",
        [
            os.environ.get("AIDRAMA_PDFTOPPM_PATH", ""),
            "/opt/homebrew/bin/pdftoppm",
            "/usr/local/bin/pdftoppm",
            "C:/Tools/poppler/Library/bin/pdftoppm.exe",
        ],
    )
    if pdftoppm:
        prefix = image_dir / safe_contract_filename(image_stem)
        for stale_image in image_dir.glob(f"{prefix.name}*.png"):
            stale_image.unlink()
        run_command([pdftoppm, "-r", "200", "-png", str(pdf_path), str(prefix)], "PDF 合同转图片失败")
        images = sorted(image_dir.glob(f"{prefix.name}-*.png"))
        if images:
            if len(images) == 1:
                target = image_dir / f"{prefix.name}.png"
                if target.exists():
                    target.unlink()
                images[0].rename(target)
                return [target]
            return images

    sips = find_command("sips") if allow_single_page_fallback else None
    if sips:
        target = image_dir / f"{safe_contract_filename(image_stem)}.png"
        run_command([sips, "-s", "format", "png", str(pdf_path), "--out", str(target)], "PDF 合同转图片失败")
        if target.exists():
            return [target]
    raise RuntimeError("无法将 PDF 合同转换为图片，请安装 poppler(pdftoppm)。")


def merge_pngs_vertically(image_paths: list[Path], target: Path) -> Path:
    if not image_paths:
        raise RuntimeError("没有可合并的合同图片。")
    if len(image_paths) == 1:
        return image_paths[0]
    try:
        from PySide6.QtGui import QColor, QImage, QPainter
    except ImportError as exception:
        raise RuntimeError("无法合并多页合同图片，请确认 PySide6 已安装。") from exception

    images = [QImage(str(path)) for path in image_paths]
    if any(image.isNull() for image in images):
        raise RuntimeError("采购合同图片读取失败，无法合并。")
    width = max(image.width() for image in images)
    height = sum(image.height() for image in images)
    merged = QImage(width, height, QImage.Format.Format_RGB32)
    merged.fill(QColor("white"))
    painter = QPainter(merged)
    try:
        y = 0
        for image in images:
            x = max((width - image.width()) // 2, 0)
            painter.drawImage(x, y, image)
            y += image.height()
    finally:
        painter.end()
    target.parent.mkdir(parents=True, exist_ok=True)
    if target.exists():
        target.unlink()
    if not merged.save(str(target), "PNG"):
        raise RuntimeError("采购合同长图保存失败。")
    for path in image_paths:
        if path != target and path.exists():
            try:
                path.unlink()
            except OSError:
                pass
    return target


def quicklook_docx_to_image(docx_path: Path, image_dir: Path, image_stem: str) -> Path | None:
    qlmanage = find_command("qlmanage")
    if not qlmanage:
        return None
    run_command([qlmanage, "-t", "-s", "1800", "-o", str(image_dir), str(docx_path)], "Quick Look 生成合同图片失败")
    generated = image_dir / f"{docx_path.name}.png"
    if not generated.exists():
        return None
    target = image_dir / f"{safe_contract_filename(image_stem)}.png"
    if generated != target:
        if target.exists():
            target.unlink()
        generated.rename(target)
    return target


def find_command(name: str, candidates: list[str] | None = None, *, prefer_candidates: bool = False) -> str | None:
    options = find_command_options(name, candidates, prefer_candidates=prefer_candidates)
    return options[0] if options else None


def find_command_options(name: str, candidates: list[str] | None = None, *, prefer_candidates: bool = False) -> list[str]:
    options: list[str] = []

    def add(value: str | None) -> None:
        value = str(value or "").strip()
        if value and value not in options:
            options.append(value)

    def add_command(value: str | None) -> None:
        value = str(value or "").strip()
        if not value:
            return
        candidate_path = Path(value)
        if candidate_path.exists():
            add(str(candidate_path))
            return
        if candidate_path.is_absolute() or "/" in value or "\\" in value:
            return
        add(shutil.which(value))

    if not prefer_candidates:
        add(shutil.which(name))
    for candidate in candidates or []:
        add_command(candidate)
    if prefer_candidates:
        add(shutil.which(name))
    return options


def run_command(command: list[str], failure_message: str) -> None:
    completed = subprocess.run(command, check=False, capture_output=True, text=True)
    if completed.returncode != 0:
        detail = (completed.stderr or completed.stdout or "").strip()
        raise RuntimeError(f"{failure_message}：{detail}" if detail else failure_message)
