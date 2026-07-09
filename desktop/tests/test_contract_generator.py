import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from aidrama_desktop.contracts import generator as contract_generator
from aidrama_desktop.contracts.generator import (
    ContractConfigStore,
    ContractRenderInput,
    all_required_contract_templates_configured,
    build_contract_template_download_path,
    contract_party_key,
    convert_pdf_to_pngs,
    convert_contract_docx_images,
    copy_contract_template,
    contract_template_key,
    format_contract_date,
    format_contract_date_compact,
    format_contract_date_no_wrap,
    format_contract_date_short,
    format_half_price,
    generate_agreement_number,
    generate_contract_start_date,
    merge_pngs_vertically,
    replace_double_date_footer_text,
    normalize_contract_docx_for_rendering,
    normalize_purchase_contract_layout,
    required_contract_template_types,
    render_contract_material_bundle,
    render_contract_docx,
)


def test_render_contract_docx_replaces_placeholders_in_paragraphs_and_tables(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("合同剧名：{{dramaTitle}}")
    doc.add_paragraph("协议编号：{{agreementNumber}}")
    doc.add_paragraph("签署日期：{{date}}")
    doc.add_paragraph("授权开始日期：{{startDate}}")
    doc.add_paragraph("半价：{{halfPrice}}")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "集数"
    table.cell(0, 1).text = "{{ episodeCount }} 集，共 {{episodeMinutes}} 分钟"
    doc.save(template)

    result = render_contract_docx(
        template,
        output,
        ContractRenderInput(
            contract_type="成本合同",
            drama_title="神医归来",
            episode_count="80",
            episode_minutes="80",
            price="5000",
            buyer="甲方",
            seller="乙方",
            sign_date="2026-06-15",
            start_date="2026-05-10",
            agreement_number="HZ-2026-06-123456",
        ),
    )

    rendered = Document(result)
    assert rendered.paragraphs[0].text == "合同剧名：神医归来"
    assert rendered.paragraphs[1].text == "协议编号：HZ-2026-06-123456"
    assert rendered.paragraphs[2].text == "签署日期：2026\u00a0年\u00a006\u00a0月\u00a015\u00a0日"
    assert rendered.paragraphs[3].text == "授权开始日期：2026\u00a0年\u00a005\u00a0月\u00a010\u00a0日"
    assert rendered.paragraphs[4].text == "半价：2500"
    assert rendered.tables[0].cell(0, 1).text == "80 集，共 80 分钟"


def test_format_half_price_removes_redundant_zeroes():
    assert format_half_price("1") == "0.5"
    assert format_half_price("2") == "1"
    assert format_half_price("0.5") == "0.25"


def test_generate_agreement_number_uses_sign_month_and_six_digit_suffix():
    number = generate_agreement_number("2026-06-15")

    assert re.fullmatch(r"HZ-2026-06-\d{6}", number)


def test_generate_contract_start_date_uses_30_to_40_days_before_sign_date():
    start_dates = {
        generate_contract_start_date("2026-06-15", f"seed-{index}")
        for index in range(20)
    }

    assert start_dates
    for start_date in start_dates:
        assert "2026-05-06" <= start_date <= "2026-05-16"


def test_format_contract_date_supports_chinese_date_input():
    assert format_contract_date("2026年6月5日") == "2026 年 06 月 05 日"


def test_format_contract_date_compact_removes_render_spacing():
    assert format_contract_date_compact("2026-06-05") == "2026年06月05日"


def test_format_contract_date_short_fits_repeated_footer_date():
    assert format_contract_date_short("2026-06-05") == "2026.06.05"


def test_format_contract_date_no_wrap_keeps_visual_spacing_without_line_breaks():
    assert format_contract_date_no_wrap("2026-06-05") == "2026\u00a0年\u00a006\u00a0月\u00a005\u00a0日"


def test_replace_double_date_footer_text_keeps_chinese_date_format():
    text = "日期：{{date}}                                            日期：{{date}}"

    replaced = replace_double_date_footer_text(text, "2026 年 06 月 23 日")

    assert replaced == "日期：2026 年 06 月 23 日                    日期：2026 年 06 月 23 日"


def test_double_date_paragraph_keeps_chinese_date_format_with_compact_spacing(tmp_path):
    template = tmp_path / "template.docx"
    output = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("日期：{{date}}                                            日期：{{date}}")
    doc.save(template)

    result = render_contract_docx(
        template,
        output,
        ContractRenderInput(
            contract_type="购买合同",
            drama_title="一五折陷阱",
            episode_count="27",
            episode_minutes="30",
            price="1",
            buyer="甲方",
            seller="乙方",
            sign_date="2026-06-23",
        ),
    )

    assert Document(result).paragraphs[0].text == (
        "日期：2026\u00a0年\u00a006\u00a0月\u00a023\u00a0日"
        "                    "
        "日期：2026\u00a0年\u00a006\u00a0月\u00a023\u00a0日"
    )


def test_render_contract_material_bundle_reuses_cached_images(tmp_path):
    templates = {}
    for contract_type in ("cost", "purchase", "rights"):
        template = tmp_path / f"{contract_type}.docx"
        doc = Document()
        doc.add_paragraph("{{contractType}} {{dramaTitle}} {{episodeCount}} {{episodeMinutes}} {{price}}")
        doc.save(template)
        templates[contract_template_key("WECHAT_VIDEO", contract_type)] = template

    converter_calls = []

    def fake_converter(docx_path: Path, image_dir: Path, image_stem: str | None):
        converter_calls.append((docx_path, image_stem))
        image_dir.mkdir(parents=True, exist_ok=True)
        image = image_dir / f"{image_stem}.png"
        image.write_bytes(b"png")
        return [image]

    def data_factory(_contract_type: str, label: str) -> ContractRenderInput:
        return ContractRenderInput(
            contract_type=label,
            drama_title="一五折陷阱",
            episode_count="27",
            episode_minutes="30",
            price="1",
            buyer="甲方",
            seller="乙方",
            sign_date="2026-06-23",
        )

    first = render_contract_material_bundle(
        templates,
        "WECHAT_VIDEO",
        tmp_path / "generated" / "task-1",
        data_factory,
        image_converter=fake_converter,
    )
    second = render_contract_material_bundle(
        templates,
        "WECHAT_VIDEO",
        tmp_path / "generated" / "task-1",
        data_factory,
        image_converter=fake_converter,
    )

    assert len(converter_calls) == 3
    assert [material.image_paths for material in second.materials] == [
        material.image_paths for material in first.materials
    ]


def test_convert_pdf_to_pngs_uses_pdftoppm_env_path(tmp_path, monkeypatch):
    pdf = tmp_path / "contract.pdf"
    pdf.write_bytes(b"%PDF")
    image_dir = tmp_path / "images"
    pdftoppm = tmp_path / "pdftoppm.exe"
    pdftoppm.write_text("", encoding="utf-8")
    monkeypatch.setenv("AIDRAMA_PDFTOPPM_PATH", str(pdftoppm))
    monkeypatch.setattr(contract_generator.shutil, "which", lambda _name: None)
    commands = []

    def fake_run_command(command: list[str], _failure_message: str) -> None:
        commands.append(command)
        prefix = Path(command[-1])
        prefix.parent.mkdir(parents=True, exist_ok=True)
        (prefix.parent / f"{prefix.name}-1.png").write_bytes(b"png")

    monkeypatch.setattr(contract_generator, "run_command", fake_run_command)

    images = convert_pdf_to_pngs(pdf, image_dir, "contract")

    assert commands[0][0] == str(pdftoppm)
    assert images == [image_dir / "contract.png"]


def test_render_contract_material_bundle_invalidates_cache_when_data_changes(tmp_path):
    templates = {}
    for contract_type in ("cost", "purchase", "rights"):
        template = tmp_path / f"{contract_type}.docx"
        doc = Document()
        doc.add_paragraph("{{price}}")
        doc.save(template)
        templates[contract_template_key("WECHAT_VIDEO", contract_type)] = template

    price = {"value": "1"}
    converter_calls = []

    def fake_converter(docx_path: Path, image_dir: Path, image_stem: str | None):
        converter_calls.append((docx_path, image_stem))
        image_dir.mkdir(parents=True, exist_ok=True)
        image = image_dir / f"{image_stem}.png"
        image.write_bytes(f"png-{len(converter_calls)}".encode("utf-8"))
        return [image]

    def data_factory(_contract_type: str, label: str) -> ContractRenderInput:
        return ContractRenderInput(
            contract_type=label,
            drama_title="一五折陷阱",
            episode_count="27",
            episode_minutes="30",
            price=price["value"],
            buyer="甲方",
            seller="乙方",
            sign_date="2026-06-23",
        )

    render_contract_material_bundle(
        templates,
        "WECHAT_VIDEO",
        tmp_path / "generated" / "task-1",
        data_factory,
        image_converter=fake_converter,
    )
    price["value"] = "2"
    render_contract_material_bundle(
        templates,
        "WECHAT_VIDEO",
        tmp_path / "generated" / "task-1",
        data_factory,
        image_converter=fake_converter,
    )

    assert len(converter_calls) == 6


def test_render_contract_material_bundle_merges_purchase_pages(tmp_path):
    from PySide6.QtGui import QColor, QImage

    templates = {}
    for contract_type in ("cost", "purchase", "rights"):
        template = tmp_path / f"{contract_type}.docx"
        doc = Document()
        doc.add_paragraph("{{price}}")
        doc.save(template)
        templates[contract_template_key("WECHAT_VIDEO", contract_type)] = template

    def fake_converter(docx_path: Path, image_dir: Path, image_stem: str | None):
        image_dir.mkdir(parents=True, exist_ok=True)
        if image_stem and image_stem.startswith("purchase-"):
            pages = []
            for index, height in enumerate((10, 12), start=1):
                image = QImage(20, height, QImage.Format.Format_RGB32)
                image.fill(QColor("white"))
                path = image_dir / f"{image_stem}-{index}.png"
                image.save(str(path), "PNG")
                pages.append(path)
            return pages
        image = image_dir / f"{image_stem or docx_path.stem}.png"
        image.write_bytes(b"png")
        return [image]

    def data_factory(_contract_type: str, label: str) -> ContractRenderInput:
        return ContractRenderInput(
            contract_type=label,
            drama_title="一五折陷阱",
            episode_count="27",
            episode_minutes="30",
            price="1",
            buyer="甲方",
            seller="乙方",
            sign_date="2026-06-23",
        )

    bundle = render_contract_material_bundle(
        templates,
        "WECHAT_VIDEO",
        tmp_path / "generated" / "task-1",
        data_factory,
        image_converter=fake_converter,
    )
    purchase = next(material for material in bundle.materials if material.contract_type == "purchase")
    merged = QImage(str(purchase.image_paths[0]))

    assert len(purchase.image_paths) == 1
    assert merged.height() == 22


def test_contract_config_store_round_trips_template_paths(tmp_path):
    store = ContractConfigStore(tmp_path / "contract-templates.json")

    store.save({
        contract_template_key("WECHAT_VIDEO", "cost"): tmp_path / "cost.docx",
        contract_template_key("WECHAT_VIDEO", "purchase"): tmp_path / "purchase.docx",
        contract_template_key("WECHAT_VIDEO", "rights"): tmp_path / "rights.docx",
    })

    assert store.load()[contract_template_key("WECHAT_VIDEO", "cost")] == tmp_path / "cost.docx"
    assert store.load()[contract_template_key("WECHAT_VIDEO", "purchase")] == tmp_path / "purchase.docx"
    assert store.load()[contract_template_key("WECHAT_VIDEO", "rights")] == tmp_path / "rights.docx"


def test_contract_config_store_round_trips_party_fields(tmp_path):
    store = ContractConfigStore(tmp_path / "contract-templates.json")

    store.save({
        contract_party_key("WECHAT_VIDEO", "buyer"): "甲方公司",
        contract_party_key("WECHAT_VIDEO", "seller"): "乙方公司",
    })

    config = store.load()

    assert config[contract_party_key("WECHAT_VIDEO", "buyer")] == "甲方公司"
    assert config[contract_party_key("WECHAT_VIDEO", "seller")] == "乙方公司"


def test_contract_config_store_migrates_legacy_template_keys(tmp_path):
    store = ContractConfigStore(tmp_path / "contract-templates.json")
    (tmp_path / "contract-templates.json").write_text(
        '{"cost": "/tmp/legacy-cost.docx", "purchase": "/tmp/legacy-purchase.docx"}',
        encoding="utf-8",
    )

    templates = store.load()

    assert templates[contract_template_key("WECHAT_VIDEO", "cost")] == Path("/tmp/legacy-cost.docx")
    assert templates[contract_template_key("WECHAT_VIDEO", "purchase")] == Path("/tmp/legacy-purchase.docx")


def test_copy_contract_template_keeps_docx_and_safe_name(tmp_path):
    source = tmp_path / "用户模板.docx"
    Document().save(source)

    target = copy_contract_template(source, tmp_path / "templates", "买剧/合同")

    assert target.exists()
    assert target.name == "买剧_合同.docx"


def test_build_contract_template_download_path_uses_selected_directory_and_safe_docx_name(tmp_path):
    target = build_contract_template_download_path(
        tmp_path,
        "wechat_video:purchase",
        {"name": "购买合同/标准版", "fileName": "系统模板"},
    )

    assert target == tmp_path / "wechat_video_purchase-购买合同_标准版.docx"


def test_merge_pngs_vertically_creates_single_long_image(tmp_path):
    from PySide6.QtGui import QColor, QImage

    first = tmp_path / "first.png"
    second = tmp_path / "second.png"
    image_1 = QImage(20, 10, QImage.Format.Format_RGB32)
    image_1.fill(QColor("white"))
    image_1.save(str(first), "PNG")
    image_2 = QImage(30, 12, QImage.Format.Format_RGB32)
    image_2.fill(QColor("white"))
    image_2.save(str(second), "PNG")

    target = merge_pngs_vertically([first, second], tmp_path / "merged.png")
    merged = QImage(str(target))

    assert target.exists()
    assert merged.width() == 30
    assert merged.height() == 22


def test_rights_contract_image_conversion_prefers_quicklook(tmp_path, monkeypatch):
    docx_path = tmp_path / "rights.docx"
    image_dir = tmp_path / "images"
    quicklook_image = image_dir / "rights.png"
    docx_path.write_bytes(b"docx")
    calls = []

    def fake_quicklook(path: Path, target_dir: Path, stem: str):
        calls.append(("quicklook", path, target_dir, stem))
        target_dir.mkdir(parents=True, exist_ok=True)
        quicklook_image.write_bytes(b"png")
        return quicklook_image

    def fake_export(path: Path, target_dir: Path, stem: str):
        calls.append(("export", path, target_dir, stem))
        return []

    monkeypatch.setattr("aidrama_desktop.contracts.generator.quicklook_docx_to_image", fake_quicklook)
    monkeypatch.setattr("aidrama_desktop.contracts.generator.export_contract_docx_images", fake_export)

    assert convert_contract_docx_images("rights", docx_path, image_dir, "rights") == [quicklook_image]
    assert calls == [("quicklook", docx_path, image_dir, "rights")]


def test_normalize_contract_docx_for_rendering_keeps_template_fonts(tmp_path):
    docx_path = tmp_path / "render.docx"
    document_xml = """<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p>
      <w:r>
        <w:rPr>
          <w:rFonts w:eastAsia="仿宋_GB2312" w:asciiTheme="minorEastAsia" w:hAnsiTheme="minorEastAsia"/>
          <w:spacing w:val="-20"/>
        </w:rPr>
        <w:t>成本配置比例情况报告</w:t>
      </w:r>
    </w:p>
  </w:body>
</w:document>"""
    with ZipFile(docx_path, "w") as archive:
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/styles.xml", "")

    assert normalize_contract_docx_for_rendering(docx_path)

    with ZipFile(docx_path) as archive:
        updated = archive.read("word/document.xml").decode("utf-8")

    assert 'w:eastAsia="仿宋_GB2312"' in updated
    assert 'w:spacing w:val="0"' in updated
    assert "minorEastAsia" in updated


def test_normalize_purchase_contract_layout_raises_stamp_without_moving_images(tmp_path):
    docx_path = tmp_path / "purchase.docx"
    original_stamp_offset = 247015
    original_seller_stamp_offset = 97155
    document_xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">
  <w:body>
    <w:p>
      <w:r>
        <w:drawing>
          <wp:anchor relativeHeight="20" behindDoc="1" allowOverlap="0">
            <wp:positionV relativeFrom="paragraph"><wp:posOffset>{original_stamp_offset}</wp:posOffset></wp:positionV>
            <wp:docPr id="1" name="图片 1" descr="甲方-盖章"/>
          </wp:anchor>
        </w:drawing>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:drawing>
          <wp:anchor relativeHeight="25" behindDoc="1" allowOverlap="0">
            <wp:positionV relativeFrom="paragraph"><wp:posOffset>{original_seller_stamp_offset}</wp:posOffset></wp:positionV>
            <wp:docPr id="3" name="图片 3" descr="乙方"/>
          </wp:anchor>
        </w:drawing>
      </w:r>
    </w:p>
    <w:p>
      <w:r>
        <w:drawing>
          <wp:anchor relativeHeight="30" behindDoc="0" allowOverlap="1">
            <wp:positionV relativeFrom="paragraph"><wp:posOffset>276225</wp:posOffset></wp:positionV>
            <wp:docPr id="2" name="图片 2" descr="签名"/>
          </wp:anchor>
        </w:drawing>
      </w:r>
    </w:p>
  </w:body>
</w:document>"""
    with ZipFile(docx_path, "w") as archive:
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/styles.xml", "")

    assert normalize_purchase_contract_layout(docx_path)

    with ZipFile(docx_path) as archive:
        updated = archive.read("word/document.xml").decode("utf-8")

    assert 'descr="甲方-盖章"' in updated
    assert 'behindDoc="0"' in updated
    assert 'allowOverlap="1"' in updated
    assert 'relativeHeight="4126"' in updated
    assert 'relativeHeight="8222"' in updated
    assert f"<wp:posOffset>{original_stamp_offset}</wp:posOffset>" in updated
    assert f"<wp:posOffset>{original_seller_stamp_offset}</wp:posOffset>" in updated
    assert not normalize_purchase_contract_layout(docx_path)
    with ZipFile(docx_path) as archive:
        assert archive.read("word/document.xml").decode("utf-8") == updated


def test_wechat_video_requires_cost_and_purchase_templates(tmp_path):
    cost = tmp_path / "cost.docx"
    purchase = tmp_path / "purchase.docx"
    rights = tmp_path / "rights.docx"
    cost.write_text("cost", encoding="utf-8")
    purchase.write_text("purchase", encoding="utf-8")
    rights.write_text("rights", encoding="utf-8")

    assert required_contract_template_types("WECHAT_VIDEO") == (
        ("cost", "成本合同"),
        ("purchase", "购买合同"),
        ("rights", "权利声明"),
    )
    assert not all_required_contract_templates_configured(
        {contract_template_key("WECHAT_VIDEO", "cost"): cost},
        "WECHAT_VIDEO",
    )
    assert all_required_contract_templates_configured(
        {
            contract_template_key("WECHAT_VIDEO", "cost"): cost,
            contract_template_key("WECHAT_VIDEO", "purchase"): purchase,
            contract_template_key("WECHAT_VIDEO", "rights"): rights,
        },
        "WECHAT_VIDEO",
    )
