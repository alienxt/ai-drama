import io
import json
import threading
import time
import urllib.error
from datetime import date
from pathlib import Path

import pytest

from aidrama_desktop.tasks.runner import (
    EpisodeMediaFile,
    TIKTOK_UPLOAD_NOT_READY_FAILURE_REASON,
    TaskRunner,
    download_episodes,
    drama_directory_name,
    episode_video_filename,
)


def drama_download_dir(tmp_path: Path, title: str = "神医归来", drama_id: str = "drama-1") -> Path:
    return tmp_path / "dramas" / "downloads" / drama_directory_name({"title": title, "dramaId": drama_id})


def drama_processed_dir(tmp_path: Path, title: str = "神医归来", drama_id: str = "drama-1") -> Path:
    return tmp_path / "dramas" / "processed" / drama_directory_name({"title": title, "dramaId": drama_id})


def last_task_result_payload(api) -> dict:
    result_calls = [payload for method, path, payload in api.calls if method == "PUT" and path.endswith("/result")]
    assert result_calls
    return result_calls[-1]


class FakeApi:
    base_url = "http://server/api"

    def __init__(self):
        self.calls = []
        self.download_token = "token-1"

    def post(self, path, payload=None):
        self.calls.append(("POST", path, payload))
        if path == "/desktop/tasks/publish-next":
            return {"id": "task-1", "dramaId": "drama-1", "mediaAccountId": "media-1"}
        if path == "/desktop/tasks/task-1/prepare":
            return {"prepared": True, "preparing": False, "failed": False, "message": "AI 素材已准备完成", "retryAfterSeconds": 0}
        if path.endswith("/pause") or path.endswith("/skip"):
            return {"id": "task-1", "status": "PENDING"}
        if path.endswith("/force-stop"):
            return {"id": "task-1", "status": "CANCELLED"}
        if path.endswith("/result"):
            return {"ok": True}
        return {}

    def get(self, path):
        self.calls.append(("GET", path, None))
        return {
            "dramaId": "drama-1",
            "title": "神医归来",
            "summary": "简介",
            "aiSummary": "AI简介...",
            "totalMinutes": 20,
            "costAmountWan": 3,
            "episodes": [
                {"episodeNo": 1, "downloadUrl": "/files/1.mp4"},
                {"episodeNo": 2, "downloadUrl": "/files/2.mp4"},
            ],
        }

    def put(self, path, payload=None):
        self.calls.append(("PUT", path, payload))
        return {}

    def download_headers(self):
        return {"Authorization": f"Bearer {self.download_token}"}


class FakeProcessor:
    def __init__(self):
        self.calls = []
        self.merge_calls = []
        self.dimensions = {}
        self.durations = {}

    def transcode_for_wechat_video(self, source: Path, target: Path, cover_path: Path | None = None) -> Path:
        self.calls.append((source, target, cover_path))
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(source.name)
        return target

    def needs_wechat_video_bitrate_transcode(self, source: Path) -> bool:
        return False

    def video_dimensions(self, source: Path):
        return self.dimensions.get(source.name)

    def video_duration_seconds(self, source: Path):
        return self.durations.get(source.name)

    def merge_videos_for_tiktok(self, sources: list[Path], target: Path) -> Path:
        self.merge_calls.append((sources, target))
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(b"merged-video")
        return target


class Strategy1Processor(FakeProcessor):
    def __init__(self):
        super().__init__()
        self.strategy1_calls = []

    def process_drama_with_strategy1(self, sources: list[Path], target_dir: Path, drama_title: str):
        self.strategy1_calls.append((sources, target_dir, drama_title))
        outputs = []
        for index, source in enumerate(sources, start=1):
            target = target_dir / f"{drama_title}-策略1第{index:03d}集.mp4"
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(f"strategy1-{source.name}")
            outputs.append(type("Segment", (), {"file": target, "source_episode_indexes": (index,)})())
        return outputs


class LowBitrateProcessor(FakeProcessor):
    def needs_wechat_video_bitrate_transcode(self, source: Path) -> bool:
        return source.name == "001.mp4"


class FailingTranscodeProcessor(LowBitrateProcessor):
    def transcode_for_wechat_video(self, source: Path, target: Path, cover_path: Path | None = None) -> Path:
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("partial")
        raise RuntimeError("Conversion failed!")


class FakePublisher:
    def __init__(self):
        self.title = None
        self.files = []
        self.summary = None
        self.metadata = None

    def publish(self, media_files, title, summary=None, metadata=None):
        self.files = media_files
        self.title = title
        self.summary = summary
        self.metadata = metadata
        return "published-1"


class FakeStoryboardGenerator:
    def __init__(self):
        self.calls = []

    def generate(self, *, source_video, drama_title, episode_label, media_account, output_dir, config):
        self.calls.append(
            {
                "source_video": source_video,
                "drama_title": drama_title,
                "episode_label": episode_label,
                "media_account": media_account,
                "output_dir": output_dir,
                "config": config,
            }
        )
        screenshots_dir = output_dir / "分镜截图"
        screenshots_dir.mkdir(parents=True, exist_ok=True)
        images = [
            screenshots_dir / "分镜-001-完整工作台.png",
            screenshots_dir / "分镜-002-完整工作台.png",
        ]
        for image in images:
            image.write_bytes(b"png")
        return images


class FailingPublisher:
    def publish(self, media_files, title, summary=None, metadata=None):
        raise RuntimeError("upload failed")


class DraftPausedPublisher:
    def publish(self, media_files, title, summary=None, metadata=None):
        from aidrama_desktop.platforms.base import PlatformPublishPaused

        raise PlatformPublishPaused("剧目提审第一步表单已填好，暂未进入下一步或提交。")


class CapturingPausedPublisher(FakePublisher):
    def publish(self, media_files, title, summary=None, metadata=None):
        super().publish(media_files, title, summary=summary, metadata=metadata)
        from aidrama_desktop.platforms.base import PlatformPublishPaused

        raise PlatformPublishPaused("TK 表单已填写并停留在提交前，等待人工核验后手动提交。")


def test_publish_once_prepares_task_and_downloads_each_episode(tmp_path, monkeypatch):
    api = FakeApi()
    processor = FakeProcessor()
    publisher = FakePublisher()
    progress_events = []

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        files = []
        total = len(download_plan["episodes"])
        for index, episode in enumerate(download_plan["episodes"], start=1):
            if progress_callback:
                progress_callback(index, total, episode, 5 * 1024 * 1024, 10 * 1024 * 1024)
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(headers["Authorization"] + " " + episode["downloadUrl"])
            files.append(target)
        return files

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)

    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
        progress_callback=lambda stage, task_id, task=None: progress_events.append((stage, task_id)),
    )

    result = runner.publish_once()

    assert result == "succeeded"
    assert ("POST", "/desktop/tasks/publish-next", {"deviceId": "device-1", "asyncPreparation": True}) in api.calls
    assert ("POST", "/desktop/tasks/task-1/prepare", None) in api.calls
    assert ("GET", "/desktop/dramas/drama-1/download-plan", None) in api.calls
    assert publisher.title == "神医归来"
    assert publisher.files == [
        drama_download_dir(tmp_path) / "001.mp4",
        drama_download_dir(tmp_path) / "002.mp4",
    ]
    assert processor.calls == []
    assert ("当前短剧：神医归来", "task-1") in progress_events
    assert ("下载：神医归来 第 1/2 集 5.0/10.0 MB（50%）", "task-1") in progress_events
    assert ("发布：神医归来", "task-1") in progress_events

def test_publish_once_runs_strategy1_after_full_download_before_upload(tmp_path, monkeypatch):
    api = FakeApi()
    processor = Strategy1Processor()
    publisher = FakePublisher()

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        files = []
        for episode in download_plan["episodes"]:
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(episode["downloadUrl"])
            files.append(target)
        return files

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)

    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "succeeded"

    strategy_sources, strategy_target_dir, drama_title = processor.strategy1_calls[0]
    assert [source.name for source in strategy_sources] == ["001.mp4", "002.mp4"]
    assert strategy_target_dir == drama_download_dir(tmp_path) / "strategy1"
    assert drama_title == "神医归来"
    assert [file.parent.name for file in publisher.files] == ["strategy1", "strategy1"]
    assert [file.name for file in publisher.files] == ["神医归来-策略1第001集.mp4", "神医归来-策略1第002集.mp4"]


def test_strategy1_sources_use_parent_download_assets_for_cover_frame(tmp_path):
    processor = FakeProcessor()
    runner = TaskRunner(
        api=FakeApi(),
        processor=processor,
        publisher=FakePublisher(),
        work_dir=tmp_path,
        device_id="device-1",
    )
    download_dir = drama_download_dir(tmp_path)
    strategy_dir = download_dir / "strategy1"
    strategy_video = strategy_dir / "神医归来-策略1第001集.mp4"
    strategy_video.parent.mkdir(parents=True)
    strategy_video.write_bytes(b"video")
    cover = download_dir / "fengmian.jpg"
    cover.write_bytes(b"cover")

    assert runner._cover_file_for_sources([strategy_video]) == cover


def test_upload_retry_reads_strategy1_manifest_for_processed_cache(tmp_path):
    runner = TaskRunner(
        api=FakeApi(),
        processor=FakeProcessor(),
        publisher=FakePublisher(),
        work_dir=tmp_path,
        device_id="device-1",
    )
    download_strategy_dir = drama_download_dir(tmp_path) / "strategy1"
    processed_strategy_dir = drama_processed_dir(tmp_path) / "strategy1"
    download_strategy_dir.mkdir(parents=True)
    processed_strategy_dir.mkdir(parents=True)
    cached_file = processed_strategy_dir / "神医归来-策略1第001集.mp4"
    cached_file.write_bytes(b"processed")
    manifest = {
        "files": [
            {
                "file": cached_file.name,
                "episodeIndex": 1,
                "episode": {
                    "episodeNo": 1,
                    "title": "第1-2集",
                    "sourceEpisodeNumbers": [1, 2],
                    "sourceEpisodeRange": "1-2",
                },
                "sourceEpisodeIndexes": [1, 2],
            }
        ]
    }
    (download_strategy_dir / ".downloaded-episodes.json").write_text(json.dumps(manifest), encoding="utf-8")

    items = runner._cached_upload_items(FakeApi().get("/desktop/dramas/drama-1/download-plan"), "WECHAT_VIDEO")

    assert len(items) == 1
    assert items[0].file == cached_file
    assert items[0].source_episode_indexes == (1, 2)
    assert items[0].episode["sourceEpisodeRange"] == "1-2"


def test_publish_once_generates_storyboard_images_for_contract_upload(tmp_path, monkeypatch):
    class StoryboardApi(FakeApi):
        def get(self, path):
            self.calls.append(("GET", path, None))
            if path == "/desktop/storyboard-config":
                return {
                    "enabled": True,
                    "deepseekApiBase": "https://api.deepseek.com",
                    "deepseekApiKey": "configured-key",
                    "deepseekModel": "deepseek-v4-pro",
                    "targetShots": 15,
                    "style": "真人风格-国产都市",
                }
            if path == "/desktop/media-accounts":
                return [{"id": "media-1", "displayName": "用户1161", "externalAccountId": "wx-1"}]
            return {
                "dramaId": "drama-1",
                "title": "神医归来",
                "summary": "简介",
                "aiSummary": "AI简介...",
                "totalMinutes": 30,
                "costAmountWan": 3,
                "episodes": [
                    {"episodeNo": 1, "downloadUrl": "/files/1.mp4"},
                    {"episodeNo": 2, "downloadUrl": "/files/2.mp4"},
                    {"episodeNo": 3, "downloadUrl": "/files/3.mp4"},
                ],
            }

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        files = []
        for index, episode in enumerate(download_plan["episodes"], start=1):
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(f"episode-{index}")
            files.append(target)
        return files

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)

    api = StoryboardApi()
    publisher = FakePublisher()
    generator = FakeStoryboardGenerator()
    progress_events = []
    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
        progress_callback=lambda stage, task_id, task=None: progress_events.append((stage, task_id)),
        storyboard_generator=generator,
        storyboards_dir=tmp_path / "storyboards",
    )

    assert runner.publish_once() == "succeeded"

    assert len(generator.calls) == 1
    call = generator.calls[0]
    assert call["source_video"] == drama_download_dir(tmp_path) / "002.mp4"
    assert call["drama_title"] == "神医归来"
    assert call["episode_label"] == "#2集"
    assert call["media_account"] == "用户1161"
    assert call["config"].enabled is True
    images = [
        tmp_path / "storyboards" / "generated" / "task-1" / "episode-2" / "分镜截图" / "分镜-001-完整工作台.png",
        tmp_path / "storyboards" / "generated" / "task-1" / "episode-2" / "分镜截图" / "分镜-002-完整工作台.png",
    ]
    assert publisher.metadata["storyboardImages"] == images
    assert publisher.metadata["buyDramaContractImages"] == images
    assert ("生成分镜图：神医归来 #2集", "task-1") in progress_events
    assert ("分镜图已生成：神医归来 #2集（2 张）", "task-1") in progress_events


def test_publish_once_copies_download_assets_to_processed_dir(tmp_path, monkeypatch):
    api = FakeApi()
    progress_events = []
    assets = {
        "fengmian.jpg": b"cover",
        "video-cover.jpg": b"video-cover",
        "fengmian-en.jpg": b"cover-en",
        "video-cover-en.jpg": b"video-cover-en",
        "tiktok-cover-en.jpg": b"tiktok-cover-en",
        "meta.json": b'{"title":"asset metadata"}',
    }

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        target_dir.mkdir(parents=True, exist_ok=True)
        for filename, body in assets.items():
            (target_dir / filename).write_bytes(body)
        target = target_dir / "001.mp4"
        target.write_bytes(b"video")
        return [target]

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)

    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=FakePublisher(),
        work_dir=tmp_path,
        device_id="device-1",
        progress_callback=lambda stage, task_id, task=None: progress_events.append((stage, task_id)),
    )

    assert runner.publish_once() == "succeeded"

    processed_dir = drama_processed_dir(tmp_path)
    for filename, body in assets.items():
        assert (processed_dir / filename).read_bytes() == body
    assert ("资料已同步到处理目录：神医归来（6 个）", "task-1") in progress_events


def test_publish_once_waits_for_async_preparation_before_download(tmp_path, monkeypatch):
    api = FakeApi()
    prepare_calls = []

    def post(path, payload=None):
        api.calls.append(("POST", path, payload))
        if path == "/desktop/tasks/publish-next":
            return {"id": "task-1", "dramaId": "drama-1", "mediaAccountId": "media-1"}
        if path == "/desktop/tasks/task-1/prepare":
            prepare_calls.append(path)
            if len(prepare_calls) == 1:
                return {
                    "prepared": False,
                    "preparing": True,
                    "failed": False,
                    "message": "AI 素材准备中，请稍候",
                    "retryAfterSeconds": 1,
                }
            return {"prepared": True, "preparing": False, "failed": False, "message": "AI 素材已准备完成", "retryAfterSeconds": 0}
        if path.endswith("/result"):
            return {"ok": True}
        return {}

    api.post = post
    monkeypatch.setattr("aidrama_desktop.tasks.runner.time.sleep", lambda _seconds: None)
    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", lambda *args, **kwargs: [tmp_path / "001.mp4"])
    (tmp_path / "001.mp4").write_text("video")
    progress_events = []
    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=FakePublisher(),
        work_dir=tmp_path,
        device_id="device-1",
        progress_callback=lambda stage, task_id, task=None: progress_events.append((stage, task_id)),
    )

    assert runner.publish_once() == "succeeded"

    assert len(prepare_calls) == 2
    assert ("AI 素材准备中，请稍候", "task-1") in progress_events
    assert ("AI素材准备完成", "task-1") in progress_events


def test_publish_once_notifies_claimed_media_account(tmp_path, monkeypatch):
    api = FakeApi()
    publisher = FakePublisher()
    progress_events = []

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        target = target_dir / "001.mp4"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("video")
        return [target]

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
        progress_callback=lambda stage, task_id, task=None: progress_events.append((stage, task_id, task)),
    )

    assert runner.publish_once() == "succeeded"

    assert ("任务已领取", "task-1", {"id": "task-1", "dramaId": "drama-1", "mediaAccountId": "media-1"}) in progress_events


def test_publish_once_passes_playlet_metadata_to_publisher(tmp_path, monkeypatch):
    api = FakeApi()
    publisher = FakePublisher()
    processor = FakeProcessor()

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        cover_file = target_dir / "fengmian.jpg"
        cover_en_file = target_dir / "fengmian-en.jpg"
        cover_file.parent.mkdir(parents=True, exist_ok=True)
        cover_file.write_text("cover")
        cover_en_file.write_text("cover-en")
        targets = []
        for episode in download_plan["episodes"]:
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.write_text("video")
            targets.append(target)
        return targets

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "succeeded"

    assert publisher.summary == "AI简介..."
    assert publisher.metadata["dramaId"] == "drama-1"
    assert publisher.metadata["publishTitle"] == "神医归来"
    assert publisher.metadata["coverFile"] == drama_download_dir(tmp_path) / "fengmian.jpg"
    assert publisher.metadata["coverEnFile"] == drama_download_dir(tmp_path) / "fengmian-en.jpg"
    assert publisher.metadata["totalMinutes"] == 20
    assert publisher.metadata["costAmountWan"] == 3
    assert publisher.metadata["productionCostWan"] == 3
    assert publisher.metadata["producerName"] == "乙方公司"
    assert publisher.metadata["aiContentDeclaration"] is True
    assert publisher.metadata["monetizationType"] == "IAA_AD"
    assert publisher.metadata["monetizationLabel"] == "IAA广告变现"
    assert publisher.metadata["freeEpisodeCount"] == 2
    assert publisher.metadata["episodeCount"] == 2
    assert publisher.metadata["originalEpisodeCount"] == 2
    assert publisher.metadata["episodes"] == [
        {
            "episodeNo": 1,
            "title": None,
            "file": drama_processed_dir(tmp_path) / "001.mp4",
        },
        {
            "episodeNo": 2,
            "title": None,
            "file": drama_processed_dir(tmp_path) / "002.mp4",
        },
    ]


def test_publish_once_transcodes_low_bitrate_video_before_upload(tmp_path, monkeypatch):
    api = FakeApi()
    publisher = FakePublisher()
    processor = LowBitrateProcessor()

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        files = []
        for episode in download_plan["episodes"]:
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text("video")
            files.append(target)
        return files

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "succeeded"

    processed_file = drama_processed_dir(tmp_path) / "001.mp4"
    untouched_file = drama_download_dir(tmp_path) / "002.mp4"
    assert ("PUT", "/desktop/tasks/task-1/progress", {"status": "PROCESSING", "progress": 70}) in api.calls
    assert publisher.files == [processed_file, untouched_file]
    assert processor.calls == [
        (
            drama_download_dir(tmp_path) / "001.mp4",
            processed_file,
            None,
        )
    ]
    assert publisher.metadata["episodes"][0]["file"] == processed_file
    assert publisher.metadata["episodes"][1]["file"] == untouched_file
    assert publisher.metadata["episodeCount"] == 2


def test_publish_once_cleans_failed_transcode_cache_for_retry(tmp_path, monkeypatch):
    api = FakeApi()
    processor = FailingTranscodeProcessor()
    publisher = FakePublisher()
    source_file = drama_download_dir(tmp_path) / "001.mp4"

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        source_file.parent.mkdir(parents=True, exist_ok=True)
        source_file.write_text("broken-video")
        return [source_file]

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "failed"

    processed_file = drama_processed_dir(tmp_path) / "001.mp4"
    result_calls = [call for call in api.calls if call[0] == "PUT" and call[1] == "/desktop/tasks/task-1/result"]
    assert result_calls
    assert "没有可上传的剧集" in result_calls[-1][2]["failureReason"]
    assert not source_file.exists()
    assert not processed_file.exists()


def test_publish_once_skips_failed_transcode_episode_and_uploads_remaining(tmp_path, monkeypatch):
    api = FakeApi()
    processor = FailingTranscodeProcessor()
    publisher = FakePublisher()

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        files = []
        for episode in download_plan["episodes"]:
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text("video")
            files.append(target)
        return files

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "succeeded"

    assert publisher.files == [drama_download_dir(tmp_path) / "002.mp4"]
    assert publisher.metadata["episodeCount"] == 1
    assert publisher.metadata["skippedEpisodeCount"] == 1
    assert publisher.metadata["skippedEpisodeNumbers"] == [1]
    assert publisher.metadata["episodes"] == [
        {
            "episodeNo": 2,
            "title": None,
            "file": drama_download_dir(tmp_path) / "002.mp4",
        },
    ]


def test_publish_once_skips_short_video_episode_and_uploads_remaining(tmp_path, monkeypatch):
    api = FakeApi()
    processor = FakeProcessor()
    processor.durations = {"001.mp4": 29.4, "002.mp4": 31.0}
    publisher = FakePublisher()
    progress_events = []

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        files = []
        for episode in download_plan["episodes"]:
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text("video")
            files.append(target)
        return files

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
        progress_callback=lambda stage, task_id, task=None: progress_events.append((stage, task_id)),
    )

    assert runner.publish_once() == "succeeded"

    assert publisher.files == [drama_download_dir(tmp_path) / "002.mp4"]
    assert publisher.metadata["episodeCount"] == 1
    assert publisher.metadata["skippedEpisodeCount"] == 1
    assert publisher.metadata["skippedEpisodeNumbers"] == [1]
    assert any("第 1 集视频时长 29.4 秒，小于视频号要求的 30 秒" in stage for stage, _task_id in progress_events)


def test_tiktok_task_stops_before_upload_after_processing(tmp_path, monkeypatch):
    from docx import Document

    api = FakeApi()
    publisher = FakePublisher()
    processor = FakeProcessor()
    processor.durations = {"001.mp4": 20.0}
    template = tmp_path / "tiktok_purchase.docx"
    document = Document()
    document.add_paragraph("{{contractType}} {{dramaTitle}} {{episodeCount}}")
    document.save(template)

    def post(path, payload=None):
        api.calls.append(("POST", path, payload))
        if path == "/desktop/tasks/publish-next":
            return {"id": "task-1", "dramaId": "drama-1", "mediaAccountId": "media-tk", "platform": "TIKTOK"}
        if path == "/desktop/tasks/task-1/prepare":
            return {"prepared": True, "preparing": False, "failed": False, "message": "AI 素材已准备完成", "retryAfterSeconds": 0}
        return {}

    def get(path):
        api.calls.append(("GET", path, None))
        return {
            "dramaId": "drama-1",
            "title": "中文剧名",
            "aiTitle": "中文AI剧名",
            "aiTitleEn": "English AI Title",
            "summary": "中文简介",
            "aiSummary": "中文AI简介",
            "aiSummaryEn": "English AI summary.",
            "episodes": [{"episodeNo": 1, "downloadUrl": "/files/1.mp4"}],
        }

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        target = target_dir / "001.mp4"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(b"0" * (6 * 1024 * 1024))
        return [target]

    def fake_image_converter(docx_path: Path, image_dir: Path, image_stem: str | None = None):
        image_dir.mkdir(parents=True, exist_ok=True)
        image = image_dir / f"{image_stem or docx_path.stem}.png"
        image.write_bytes(b"png")
        return [image]

    api.post = post
    api.get = get
    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
        contract_templates={"tiktok:purchase": template},
        contracts_dir=tmp_path / "contracts",
        contract_image_converter=fake_image_converter,
    )

    assert runner.publish_once() == "failed"

    assert publisher.title is None
    assert publisher.files == []
    result_payload = last_task_result_payload(api)
    assert result_payload == {
        "success": False,
        "platformPublishId": None,
        "failureReason": TIKTOK_UPLOAD_NOT_READY_FAILURE_REASON,
    }
    assert ("PUT", "/desktop/tasks/task-1/progress", {"status": "UPLOADING", "progress": 75}) in api.calls


def test_tiktok_task_reuses_regular_processed_files_when_they_satisfy_upload_rules(tmp_path, monkeypatch):
    api = FakeApi()
    publisher = FakePublisher()

    class DownloadOnlyTranscodeProcessor(FakeProcessor):
        def needs_wechat_video_bitrate_transcode(self, source: Path) -> bool:
            return "downloads" in source.parts

    processor = DownloadOnlyTranscodeProcessor()

    def post(path, payload=None):
        api.calls.append(("POST", path, payload))
        if path == "/desktop/tasks/publish-next":
            return {"id": "task-1", "dramaId": "drama-1", "mediaAccountId": "media-tk", "platform": "TIKTOK"}
        if path == "/desktop/tasks/task-1/prepare":
            return {"prepared": True, "preparing": False, "failed": False, "message": "AI 素材已准备完成", "retryAfterSeconds": 0}
        return {}

    def get(path):
        api.calls.append(("GET", path, None))
        return {
            "dramaId": "drama-1",
            "title": "复用短剧",
            "aiTitleEn": "Reusable Drama",
            "aiSummaryEn": "A ready processed drama.",
            "episodes": [{"episodeNo": 1, "downloadUrl": "/files/1.mp4"}],
        }

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        target = target_dir / "001.mp4"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(b"source-video")
        return [target]

    processed_file = drama_processed_dir(tmp_path, "复用短剧") / "001.mp4"
    processed_file.parent.mkdir(parents=True, exist_ok=True)
    processed_file.write_bytes(b"processed-video")

    api.post = post
    api.get = get
    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    monkeypatch.setattr("aidrama_desktop.tasks.runner.TIKTOK_MIN_VIDEO_SIZE_BYTES", 1)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "failed"

    assert publisher.files == []
    assert processed_file.exists()
    assert last_task_result_payload(api)["failureReason"] == TIKTOK_UPLOAD_NOT_READY_FAILURE_REASON
    assert processor.calls == []
    assert processor.merge_calls == []
    assert not (drama_processed_dir(tmp_path, "复用短剧") / "TK").exists()


def test_tiktok_task_merges_episode_pairs_when_upload_limit_exceeded(tmp_path, monkeypatch):
    from docx import Document

    api = FakeApi()
    publisher = FakePublisher()
    processor = FakeProcessor()
    template = tmp_path / "tiktok_purchase.docx"
    document = Document()
    document.add_paragraph("{{episodeCount}} {{episodeMinutes}}")
    document.save(template)

    def post(path, payload=None):
        api.calls.append(("POST", path, payload))
        if path == "/desktop/tasks/publish-next":
            return {"id": "task-1", "dramaId": "drama-1", "mediaAccountId": "media-tk", "platform": "TIKTOK"}
        if path == "/desktop/tasks/task-1/prepare":
            return {"prepared": True, "preparing": False, "failed": False, "message": "AI 素材已准备完成", "retryAfterSeconds": 0}
        return {}

    def get(path):
        api.calls.append(("GET", path, None))
        return {
            "dramaId": "drama-1",
            "title": "轮回苟到天荒",
            "aiTitleEn": "Endless Reincarnation",
            "aiSummaryEn": "A fantasy drama.",
            "episodeCount": 168,
            "totalMinutes": 168,
            "episodes": [
                {"episodeNo": episode_no, "downloadUrl": f"/files/{episode_no}.mp4"}
                for episode_no in range(1, 169)
            ],
        }

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        files = []
        for episode in download_plan["episodes"]:
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_bytes(b"video")
            files.append(target)
        return files

    def fake_image_converter(docx_path: Path, image_dir: Path, image_stem: str | None = None):
        image_dir.mkdir(parents=True, exist_ok=True)
        image = image_dir / f"{image_stem or docx_path.stem}.png"
        image.write_bytes(b"png")
        return [image]

    api.post = post
    api.get = get
    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    monkeypatch.setattr("aidrama_desktop.tasks.runner.TIKTOK_MIN_VIDEO_SIZE_BYTES", 1)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
        contract_templates={"tiktok:purchase": template},
        contracts_dir=tmp_path / "contracts",
        contract_image_converter=fake_image_converter,
    )

    assert runner.publish_once() == "failed"

    tiktok_dir = drama_download_dir(tmp_path, "轮回苟到天荒") / "TK"
    assert len(processor.merge_calls) == 84
    merged_files = sorted(tiktok_dir.glob("*.mp4"))
    assert len(merged_files) == 84
    assert [path.name for path in merged_files[:2]] == [
        "轮回苟到天荒-TK第001集.mp4",
        "轮回苟到天荒-TK第002集.mp4",
    ]
    assert processor.merge_calls[0][0] == [
        drama_download_dir(tmp_path, "轮回苟到天荒") / "001.mp4",
        drama_download_dir(tmp_path, "轮回苟到天荒") / "002.mp4",
    ]
    assert processor.merge_calls[-1][0] == [
        drama_download_dir(tmp_path, "轮回苟到天荒") / "167.mp4",
        drama_download_dir(tmp_path, "轮回苟到天荒") / "168.mp4",
    ]
    assert publisher.files == []
    assert last_task_result_payload(api)["failureReason"] == TIKTOK_UPLOAD_NOT_READY_FAILURE_REASON


def test_tiktok_task_merges_before_transcoding_when_upload_limit_exceeded(tmp_path, monkeypatch):
    api = FakeApi()
    publisher = FakePublisher()

    class AlwaysTranscodeProcessor(FakeProcessor):
        def needs_wechat_video_bitrate_transcode(self, source: Path) -> bool:
            return True

    processor = AlwaysTranscodeProcessor()

    def post(path, payload=None):
        api.calls.append(("POST", path, payload))
        if path == "/desktop/tasks/publish-next":
            return {"id": "task-1", "dramaId": "drama-1", "mediaAccountId": "media-tk", "platform": "TIKTOK"}
        if path == "/desktop/tasks/task-1/prepare":
            return {"prepared": True, "preparing": False, "failed": False, "message": "AI 素材已准备完成", "retryAfterSeconds": 0}
        return {}

    def get(path):
        api.calls.append(("GET", path, None))
        return {
            "dramaId": "drama-1",
            "title": "轮回苟到天荒",
            "episodeCount": 168,
            "totalMinutes": 168,
            "episodes": [
                {"episodeNo": episode_no, "downloadUrl": f"/files/{episode_no}.mp4"}
                for episode_no in range(1, 169)
            ],
        }

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        files = []
        for episode in download_plan["episodes"]:
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_bytes(b"video")
            files.append(target)
        return files

    api.post = post
    api.get = get
    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    monkeypatch.setattr("aidrama_desktop.tasks.runner.TIKTOK_MIN_VIDEO_SIZE_BYTES", 1)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "failed"

    merged_dir = drama_download_dir(tmp_path, "轮回苟到天荒") / "TK"
    processed_merged_dir = drama_processed_dir(tmp_path, "轮回苟到天荒") / "TK"
    assert len(processor.merge_calls) == 84
    assert len(processor.calls) == 84
    assert all(call[0].parent == merged_dir for call in processor.calls)
    assert all(call[1].parent == processed_merged_dir for call in processor.calls)
    assert len(list(processed_merged_dir.glob("*.mp4"))) == 84
    assert publisher.files == []
    assert last_task_result_payload(api)["failureReason"] == TIKTOK_UPLOAD_NOT_READY_FAILURE_REASON


def test_tiktok_task_merges_last_three_when_episode_count_is_odd(tmp_path, monkeypatch):
    processor = FakeProcessor()
    runner = TaskRunner(
        api=FakeApi(),
        processor=processor,
        publisher=FakePublisher(),
        work_dir=tmp_path,
        device_id="device-1",
    )
    source_dir = tmp_path / "dramas" / "downloads" / "drama-odd"
    items = []
    for episode_no in range(1, 122):
        target = source_dir / f"{episode_no:03d}.mp4"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(b"video")
        items.append(EpisodeMediaFile({"episodeNo": episode_no}, episode_no, target))

    monkeypatch.setattr("aidrama_desktop.tasks.runner.TIKTOK_MIN_VIDEO_SIZE_BYTES", 1)
    merged_items = runner._filter_upload_items_for_platform(items, "task-1", "奇数短剧", "TIKTOK")

    assert len(merged_items) == 60
    assert len(processor.merge_calls) == 60
    assert processor.merge_calls[-1][0] == [
        source_dir / "119.mp4",
        source_dir / "120.mp4",
        source_dir / "121.mp4",
    ]
    assert merged_items[-1].episode["episodeNo"] == 60
    assert merged_items[-1].episode["sourceEpisodeNumbers"] == [119, 120, 121]
    assert merged_items[-1].source_episode_indexes == (119, 120, 121)


def test_tiktok_upload_retry_uses_cached_merged_videos_without_processing(tmp_path, monkeypatch):
    api = FakeApi()
    publisher = CapturingPausedPublisher()
    processor = LowBitrateProcessor()
    cached_dir = tmp_path / "dramas" / "processed" / "drama-1" / "TK"
    cached_dir.mkdir(parents=True)
    cached_files = []
    for episode_no in range(1, 85):
        target = cached_dir / f"轮回苟到天荒-TK第{episode_no:03d}集.mp4"
        target.write_bytes(b"merged")
        cached_files.append(target)

    def get(path):
        api.calls.append(("GET", path, None))
        return {
            "dramaId": "drama-1",
            "title": "轮回苟到天荒",
            "aiTitleEn": "Endless Reincarnation",
            "aiSummaryEn": "A fantasy drama.",
            "episodeCount": 168,
            "totalMinutes": 168,
            "episodes": [
                {"episodeNo": episode_no, "downloadUrl": f"/files/{episode_no}.mp4"}
                for episode_no in range(1, 169)
            ],
        }

    api.get = get
    monkeypatch.setattr("aidrama_desktop.tasks.runner.TIKTOK_MIN_VIDEO_SIZE_BYTES", 1)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    result = runner.execute_task_from_upload_cache(
        {"id": "task-1", "dramaId": "drama-1", "mediaAccountId": "media-tk", "platform": "TIKTOK"}
    )

    assert result == "ready-for-review"
    assert publisher.files == cached_files
    assert publisher.title == "Endless Reincarnation"
    assert publisher.summary == "A fantasy drama."
    assert publisher.metadata["episodeCount"] == 84
    assert publisher.metadata["totalMinutes"] == 84
    assert publisher.metadata["originalEpisodeCount"] == 168
    assert publisher.metadata["skippedEpisodeCount"] == 0
    assert publisher.metadata["episodes"][0]["sourceEpisodeNumbers"] == [1, 2]
    assert publisher.metadata["episodes"][-1]["sourceEpisodeNumbers"] == [167, 168]
    assert processor.calls == []
    assert processor.merge_calls == []
    result_payload = last_task_result_payload(api)
    assert result_payload["success"] is False
    assert "等待人工核验" in result_payload["failureReason"]


def test_publish_once_fails_when_all_video_episodes_are_too_short(tmp_path, monkeypatch):
    api = FakeApi()
    processor = FakeProcessor()
    processor.durations = {"001.mp4": 12.0, "002.mp4": 29.9}
    publisher = FakePublisher()

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        files = []
        for episode in download_plan["episodes"]:
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text("video")
            files.append(target)
        return files

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "failed"

    result_calls = [call for call in api.calls if call[0] == "PUT" and call[1] == "/desktop/tasks/task-1/result"]
    assert result_calls
    failure_reason = result_calls[-1][2]["failureReason"]
    assert "没有可上传的剧集" in failure_reason
    assert "第 2 集视频时长 29.9 秒，小于视频号要求的 30 秒" in failure_reason
    assert publisher.files == []


def test_publish_once_transcodes_low_resolution_video(tmp_path, monkeypatch):
    api = FakeApi()
    processor = FakeProcessor()
    processor.dimensions = {"001.mp4": (540, 960), "002.mp4": (720, 1280)}
    publisher = FakePublisher()

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        files = []
        for episode in download_plan["episodes"]:
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text("video")
            files.append(target)
        return files

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "succeeded"

    source_1 = drama_download_dir(tmp_path) / "001.mp4"
    source_2 = drama_download_dir(tmp_path) / "002.mp4"
    processed_1 = drama_processed_dir(tmp_path) / "001.mp4"
    assert processor.calls == [(source_1, processed_1, None)]
    assert publisher.files == [processed_1, source_2]
    marker = processed_1.with_name("001.mp4.aidrama.json")
    signature = json.loads(marker.read_text(encoding="utf-8"))
    assert signature["minWidth"] == 720
    assert signature["minHeight"] == 1280


def test_publish_once_adds_cover_frame_to_processed_videos(tmp_path, monkeypatch):
    api = FakeApi()
    publisher = FakePublisher()
    processor = LowBitrateProcessor()
    processor.dimensions = {"001.mp4": (720, 1280)}

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        cover_file = target_dir / "fengmian.jpg"
        cover_file.parent.mkdir(parents=True, exist_ok=True)
        cover_file.write_text("cover")
        files = []
        for episode in download_plan["episodes"]:
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.write_text("video")
            files.append(target)
        return files

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "succeeded"

    cover_file = drama_download_dir(tmp_path) / "fengmian.jpg"
    processed_1 = drama_processed_dir(tmp_path) / "001.mp4"
    processed_2 = drama_processed_dir(tmp_path) / "002.mp4"
    assert publisher.files == [processed_1, processed_2]
    assert processor.calls == [
        (
            drama_download_dir(tmp_path) / "001.mp4",
            processed_1,
            cover_file,
        ),
        (
            drama_download_dir(tmp_path) / "002.mp4",
            processed_2,
            cover_file,
        ),
    ]
    marker = processed_1.with_name("001.mp4.aidrama.json")
    assert json.loads(marker.read_text(encoding="utf-8"))["cover"] == str(cover_file)
    assert json.loads(marker.read_text(encoding="utf-8"))["source"] == str(
        drama_download_dir(tmp_path) / "001.mp4"
    )


def test_publish_once_uses_video_cover_for_horizontal_videos(tmp_path, monkeypatch):
    api = FakeApi()
    publisher = FakePublisher()
    processor = LowBitrateProcessor()
    processor.dimensions = {"001.mp4": (1280, 720)}

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        cover_file = target_dir / "fengmian.jpg"
        video_cover_file = target_dir / "video-cover.jpg"
        cover_file.parent.mkdir(parents=True, exist_ok=True)
        cover_file.write_text("poster-cover")
        video_cover_file.write_text("video-cover")
        files = []
        for episode in download_plan["episodes"]:
            target = target_dir / f"{episode['episodeNo']:03d}.mp4"
            target.write_text("video")
            files.append(target)
        return files

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "succeeded"

    video_cover_file = drama_download_dir(tmp_path) / "video-cover.jpg"
    processed_1 = drama_processed_dir(tmp_path) / "001.mp4"
    assert processor.calls[0] == (
        drama_download_dir(tmp_path) / "001.mp4",
        processed_1,
        video_cover_file,
    )
    marker = processed_1.with_name("001.mp4.aidrama.json")
    assert json.loads(marker.read_text(encoding="utf-8"))["cover"] == str(video_cover_file)


def test_publish_metadata_uses_stable_free_episode_count():
    assert TaskRunner._free_episode_count({"freeEpisodeCount": 6}, 49) == 6
    assert TaskRunner._free_episode_count({}, 49) == 10
    assert TaskRunner._free_episode_count({}, 80) == 16
    assert TaskRunner._free_episode_count({}, 120) == 20
    assert TaskRunner._free_episode_count({}, 2) == 2


def test_publish_once_generates_contract_materials_before_upload(tmp_path, monkeypatch):
    from docx import Document

    api = FakeApi()
    publisher = FakePublisher()
    templates = {}
    for contract_type in ("cost", "purchase", "rights"):
        template = tmp_path / f"{contract_type}.docx"
        document = Document()
        document.add_paragraph(
            "{{contractType}} {{agreementNumber}} {{dramaTitle}} {{episodeCount}} {{episodeMinutes}} {{price}} {{halfPrice}}"
        )
        document.add_paragraph("签署日期：{{date}}")
        document.add_paragraph("授权开始日期：{{startDate}}")
        document.save(template)
        templates[f"wechat_video:{contract_type}"] = template

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        target = target_dir / "001.mp4"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("video")
        return [target]

    def fake_image_converter(docx_path: Path, image_dir: Path, image_stem: str | None = None):
        image_dir.mkdir(parents=True, exist_ok=True)
        image = image_dir / f"{image_stem or docx_path.stem}.png"
        image.write_bytes(b"png")
        return [image]

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)

    class FixedDate(date):
        @classmethod
        def today(cls):
            return cls(2026, 7, 7)

    monkeypatch.setattr("aidrama_desktop.tasks.runner.date", FixedDate)
    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
        contract_templates=templates,
        contracts_dir=tmp_path / "contracts",
        contract_buyer="甲方",
        contract_seller="乙方",
        contract_image_converter=fake_image_converter,
    )

    assert runner.publish_once() == "succeeded"

    assert publisher.metadata["purchaseContractDocx"].exists()
    assert publisher.metadata["costContractDocx"].exists()
    assert publisher.metadata["rightsStatementDocx"].exists()
    assert len(publisher.metadata["purchaseContractImages"]) == 1
    assert len(publisher.metadata["rightsStatementImages"]) == 1
    assert len(publisher.metadata["buyDramaContractImages"]) == 2
    assert len(publisher.metadata["costConfigReportImages"]) == 1
    assert publisher.metadata["buyDramaContractImages"][0].exists()
    assert publisher.metadata["buyDramaContractImages"][1].exists()
    assert publisher.metadata["costConfigReportImages"][0].exists()
    purchase_number = Document(publisher.metadata["purchaseContractDocx"]).paragraphs[0].text.split()[1]
    cost_number = Document(publisher.metadata["costContractDocx"]).paragraphs[0].text.split()[1]
    rights_number = Document(publisher.metadata["rightsStatementDocx"]).paragraphs[0].text.split()[1]
    purchase_paragraphs = Document(publisher.metadata["purchaseContractDocx"]).paragraphs
    assert purchase_number.startswith("HZ-")
    assert purchase_number.startswith("HZ-2026-07-")
    assert cost_number == purchase_number
    assert rights_number == purchase_number
    assert purchase_paragraphs[1].text == "签署日期：2026年07月06日"
    assert "2026年05月" in purchase_paragraphs[2].text or "2026年06月" in purchase_paragraphs[2].text


def test_publish_once_generates_contracts_with_successful_episode_count(tmp_path, monkeypatch):
    from docx import Document

    api = FakeApi()
    publisher = FakePublisher()
    download_plan = {
        "dramaId": "drama-1",
        "title": "跳集短剧",
        "summary": "简介",
        "totalMinutes": 30,
        "costAmountWan": 3,
        "episodes": [
            {"episodeNo": 1, "downloadUrl": "/files/1.mp4"},
            {"episodeNo": 2, "downloadUrl": "/files/2.mp4"},
            {"episodeNo": 3, "downloadUrl": "/files/3.mp4"},
        ],
    }

    def get(path):
        api.calls.append(("GET", path, None))
        return download_plan

    api.get = get
    templates = {}
    for contract_type in ("cost", "purchase", "rights"):
        template = tmp_path / f"{contract_type}.docx"
        document = Document()
        document.add_paragraph("{{episodeCount}} {{episodeMinutes}}")
        document.save(template)
        templates[f"wechat_video:{contract_type}"] = template

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        files = []
        for episode in (download_plan["episodes"][0], download_plan["episodes"][2]):
            target = target_dir / f"跳集短剧-第{episode['episodeNo']}集.mp4"
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text("video")
            files.append(target)
        return files

    def fake_image_converter(docx_path: Path, image_dir: Path, image_stem: str | None = None):
        image_dir.mkdir(parents=True, exist_ok=True)
        image = image_dir / f"{image_stem or docx_path.stem}.png"
        image.write_bytes(b"png")
        return [image]

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
        contract_templates=templates,
        contracts_dir=tmp_path / "contracts",
        contract_image_converter=fake_image_converter,
    )

    assert runner.publish_once() == "succeeded"

    assert publisher.metadata["episodeCount"] == 2
    assert publisher.metadata["totalMinutes"] == 20
    assert publisher.metadata["skippedEpisodeNumbers"] == [2]
    assert Document(publisher.metadata["costContractDocx"]).paragraphs[0].text == "2 20"


def test_publish_once_uses_media_account_specific_publisher(tmp_path, monkeypatch):
    api = FakeApi()
    publishers = {}

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        target = target_dir / "001.mp4"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("video")
        return [target]

    def publisher_factory(media_account_id):
        publisher = FakePublisher()
        publishers[media_account_id] = publisher
        return publisher

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=FakePublisher(),
        work_dir=tmp_path,
        device_id="device-1",
        publisher_factory=publisher_factory,
    )

    assert runner.publish_once() == "succeeded"
    assert "media-1" in publishers
    assert publishers["media-1"].title == "神医归来"


def test_publish_once_marks_task_failed_when_upload_fails(tmp_path, monkeypatch):
    api = FakeApi()

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        target = target_dir / "001.mp4"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("video")
        return [target]

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=FailingPublisher(),
        work_dir=tmp_path,
        device_id="device-1",
    )

    result = runner.publish_once()

    assert result == "failed"
    assert (
        "PUT",
        "/desktop/tasks/task-1/result",
        {"success": False, "platformPublishId": None, "failureReason": "upload failed"},
    ) in api.calls


def test_execute_task_from_upload_cache_skips_download_and_processing(tmp_path, monkeypatch):
    api = FakeApi()
    processor = LowBitrateProcessor()
    publisher = FakePublisher()
    cached_files = []
    for episode_no in (1, 2):
        target = drama_processed_dir(tmp_path) / f"{episode_no:03d}.mp4"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("processed-video")
        cached_files.append(target)

    def fail_download(*_args, **_kwargs):
        raise AssertionError("download should not be called when retrying from upload cache")

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fail_download)
    runner = TaskRunner(
        api=api,
        processor=processor,
        publisher=publisher,
        work_dir=tmp_path,
        device_id="device-1",
    )

    result = runner.execute_task_from_upload_cache(
        {"id": "task-1", "dramaId": "drama-1", "mediaAccountId": "media-1", "platform": "WECHAT_VIDEO"}
    )

    assert result == "succeeded"
    assert publisher.files == cached_files
    assert processor.calls == []
    assert processor.merge_calls == []
    assert ("PUT", "/desktop/tasks/task-1/progress", {"status": "UPLOADING", "progress": 75}) in api.calls
    assert last_task_result_payload(api) == {
        "success": True,
        "platformPublishId": "published-1",
        "failureReason": None,
    }


def test_publish_once_marks_task_cancelled_when_download_is_stopped(tmp_path, monkeypatch):
    api = FakeApi()

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        from aidrama_desktop.tasks.runner import TaskCancelled

        raise TaskCancelled("用户停止下载")

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=FakePublisher(),
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "cancelled"
    assert (
        "POST",
        "/desktop/tasks/task-1/force-stop",
        None,
    ) in api.calls


def test_publish_once_pauses_task_without_cancelling(tmp_path, monkeypatch):
    api = FakeApi()

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        from aidrama_desktop.tasks.runner import TaskPaused

        raise TaskPaused("用户暂停任务")

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=FakePublisher(),
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "paused"
    assert ("POST", "/desktop/tasks/task-1/pause", {"deviceId": "device-1"}) in api.calls
    assert not any(call[1] == "/desktop/tasks/task-1/progress" and call[2]["status"] == "CANCELLED" for call in api.calls)


def test_publish_once_fails_upload_when_playlet_first_step_is_ready(tmp_path, monkeypatch):
    api = FakeApi()

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        target = target_dir / "001.mp4"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("video")
        return [target]

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=DraftPausedPublisher(),
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "failed"
    assert (
        "PUT",
        "/desktop/tasks/task-1/result",
        {
            "success": False,
            "platformPublishId": None,
            "failureReason": "剧目提审第一步表单已填好，暂未进入下一步或提交。",
        },
    ) in api.calls
    assert ("POST", "/desktop/tasks/task-1/pause", {"deviceId": "device-1"}) not in api.calls


def test_publish_once_fails_when_playlet_stops_before_final_submit(tmp_path, monkeypatch):
    api = FakeApi()

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        target = target_dir / "001.mp4"
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text("video")
        return [target]

    class StoppedBeforeSubmitPublisher:
        def publish(self, media_files, title, summary=None, metadata=None):
            from aidrama_desktop.platforms.base import PlatformPublishPaused

            raise PlatformPublishPaused("剧目提审信息确认页已打开，已停留在最后一步，等待手动提交审核。")

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=StoppedBeforeSubmitPublisher(),
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "failed"
    assert (
        "PUT",
        "/desktop/tasks/task-1/result",
        {
            "success": False,
            "platformPublishId": None,
            "failureReason": "剧目提审信息确认页已打开，已停留在最后一步，等待手动提交审核。",
        },
    ) in api.calls


def test_publish_once_skips_task_without_cancelling(tmp_path, monkeypatch):
    api = FakeApi()

    def fake_download(download_plan, target_dir, base_url, headers=None, progress_callback=None, should_stop=None, should_pause=None, should_skip=None, max_concurrent_downloads=6):
        from aidrama_desktop.tasks.runner import TaskSkipped

        raise TaskSkipped("用户跳过任务")

    monkeypatch.setattr("aidrama_desktop.tasks.runner.download_episodes", fake_download)
    runner = TaskRunner(
        api=api,
        processor=FakeProcessor(),
        publisher=FakePublisher(),
        work_dir=tmp_path,
        device_id="device-1",
    )

    assert runner.publish_once() == "skipped"
    assert ("POST", "/desktop/tasks/task-1/skip", {"deviceId": "device-1"}) in api.calls
    assert not any(call[1] == "/desktop/tasks/task-1/progress" and call[2]["status"] == "CANCELLED" for call in api.calls)


def test_download_episodes_writes_cover_and_metadata(tmp_path, monkeypatch):
    opened_urls = []

    class FakeResponse:
        def __init__(self, body: bytes):
            self.body = body
            self.offset = 0
            self.headers = {"Content-Length": str(len(body))}

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self, size: int) -> bytes:
            if self.offset >= len(self.body):
                return b""
            chunk = self.body[self.offset : self.offset + size]
            self.offset += len(chunk)
            return chunk

    def fake_urlopen(request):
        opened_urls.append(request.full_url)
        if request.full_url.endswith("/uploads/covers/drama.jpg"):
            return FakeResponse(b"cover")
        if request.full_url.endswith("/uploads/covers/video.jpg"):
            return FakeResponse(b"video-cover")
        if request.full_url.endswith("/uploads/covers/drama-en.jpg"):
            return FakeResponse(b"cover-en")
        if request.full_url.endswith("/uploads/covers/video-en.jpg"):
            return FakeResponse(b"video-cover-en")
        return FakeResponse(b"video")

    monkeypatch.setattr("urllib.request.urlopen", fake_urlopen)
    download_plan = {
        "dramaId": "drama-1",
        "title": "神医归来",
        "aiTitle": "神医归来AI",
        "summary": "简介",
        "aiSummary": "AI简介...",
        "effectiveCoverUrl": "/uploads/covers/drama.jpg",
        "aiVideoCoverUrl": "/uploads/covers/video.jpg",
        "aiTitleEn": "English Title",
        "aiSummaryEn": "English summary.",
        "aiCoverEnUrl": "/uploads/covers/drama-en.jpg",
        "aiVideoCoverEnUrl": "/uploads/covers/video-en.jpg",
        "rating": 5,
        "categoryIds": ["urban"],
        "episodes": [
            {"episodeNo": 1, "title": "第一集", "sourcePath": "/pan/001.mp4", "downloadUrl": "/files/1.mp4"},
        ],
    }

    files = download_episodes(download_plan, tmp_path / "drama-1", "http://server/api")

    episode_file = tmp_path / "drama-1" / "神医归来AI-第1集.mp4"
    assert files == [episode_file]
    assert (tmp_path / "drama-1" / "fengmian.jpg").read_bytes() == b"cover"
    assert (tmp_path / "drama-1" / "video-cover.jpg").read_bytes() == b"video-cover"
    assert (tmp_path / "drama-1" / "fengmian-en.jpg").read_bytes() == b"cover-en"
    assert (tmp_path / "drama-1" / "video-cover-en.jpg").read_bytes() == b"video-cover-en"
    assert episode_file.read_bytes() == b"video"
    metadata = json.loads((tmp_path / "drama-1" / "meta.json").read_text(encoding="utf-8"))
    assert metadata["title"] == "神医归来"
    assert metadata["aiTitleEn"] == "English Title"
    assert metadata["publishTitle"] == "神医归来AI"
    assert metadata["summary"] == "AI简介..."
    assert metadata["aiSummary"] == "AI简介..."
    assert metadata["aiSummaryEn"] == "English summary."
    assert metadata["originalSummary"] == "简介"
    assert metadata["coverFile"] == "fengmian.jpg"
    assert metadata["coverEnFile"] == "fengmian-en.jpg"
    assert metadata["videoCoverFile"] == "video-cover.jpg"
    assert metadata["videoCoverEnFile"] == "video-cover-en.jpg"
    assert metadata["videoCoverUrl"] == "/uploads/covers/video.jpg"
    assert metadata["coverEnUrl"] == "/uploads/covers/drama-en.jpg"
    assert metadata["videoCoverEnUrl"] == "/uploads/covers/video-en.jpg"
    assert metadata["episodeCount"] == 1
    assert metadata["episodes"][0]["fileName"] == "神医归来AI-第1集.mp4"
    assert metadata["episodes"][0]["size"] is None
    assert opened_urls == [
        "http://server/uploads/covers/drama.jpg",
        "http://server/uploads/covers/video.jpg",
        "http://server/uploads/covers/drama-en.jpg",
        "http://server/uploads/covers/video-en.jpg",
        "http://server/files/1.mp4",
    ]


def test_download_episodes_prepares_tiktok_cover_from_english_cover(tmp_path, monkeypatch):
    from PySide6.QtGui import QColor, QImage

    source_cover = tmp_path / "source-cover.jpg"
    image = QImage(1024, 1536, QImage.Format.Format_RGB32)
    image.fill(QColor("red"))
    image.save(str(source_cover), "JPEG")
    cover_bytes = source_cover.read_bytes()
    opened_urls = []

    class FakeResponse:
        def __init__(self, body: bytes):
            self.body = body
            self.offset = 0
            self.headers = {"Content-Length": str(len(body))}

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self, size: int) -> bytes:
            if self.offset >= len(self.body):
                return b""
            chunk = self.body[self.offset : self.offset + size]
            self.offset += len(chunk)
            return chunk

    def fake_urlopen(request):
        opened_urls.append(request.full_url)
        if request.full_url.endswith("/uploads/covers/drama-en.jpg"):
            return FakeResponse(cover_bytes)
        return FakeResponse(b"video")

    monkeypatch.setattr("urllib.request.urlopen", fake_urlopen)
    download_plan = {
        "dramaId": "drama-1",
        "title": "TK封面短剧",
        "aiCoverEnUrl": "/uploads/covers/drama-en.jpg",
        "episodes": [
            {"episodeNo": 1, "sourcePath": "/pan/001.mp4", "downloadUrl": "/files/1.mp4"},
        ],
    }

    download_episodes(download_plan, tmp_path / "drama-1", "http://server/api")

    tiktok_cover = tmp_path / "drama-1" / "tiktok-cover-en.jpg"
    assert tiktok_cover.exists()
    prepared = QImage(str(tiktok_cover))
    assert prepared.width() == 768
    assert prepared.height() == 1024
    assert tiktok_cover.stat().st_size < 10 * 1024 * 1024
    metadata = json.loads((tmp_path / "drama-1" / "meta.json").read_text(encoding="utf-8"))
    assert metadata["coverEnFile"] == "fengmian-en.jpg"
    assert metadata["tiktokCoverEnFile"] == "tiktok-cover-en.jpg"
    assert opened_urls == [
        "http://server/uploads/covers/drama-en.jpg",
        "http://server/files/1.mp4",
    ]


def test_download_episodes_skips_complete_existing_episode(tmp_path, monkeypatch):
    target_dir = tmp_path / "drama-1"
    target_dir.mkdir()
    legacy_file = target_dir / "001.mp4"
    legacy_file.write_bytes(b"already-downloaded")
    expected_file = target_dir / "续传短剧-第1集.mp4"
    opened_urls = []

    def fake_urlopen(request):
        opened_urls.append(request.full_url)
        raise AssertionError("complete episode should not be downloaded again")

    monkeypatch.setattr("urllib.request.urlopen", fake_urlopen)
    download_plan = {
        "dramaId": "drama-1",
        "title": "续传短剧",
        "episodes": [
            {
                "episodeNo": 1,
                "sourcePath": "/pan/001.mp4",
                "size": len(b"already-downloaded"),
                "downloadUrl": "/files/1.mp4",
            },
        ],
    }

    files = download_episodes(download_plan, target_dir, "http://server/api")

    assert files == [expected_file]
    assert expected_file.read_bytes() == b"already-downloaded"
    assert not legacy_file.exists()
    assert opened_urls == []


def test_download_episodes_retries_retryable_download_error(tmp_path, monkeypatch):
    opened_urls = []
    attempts = 0

    class FakeResponse:
        headers = {"Content-Length": "5"}

        def __enter__(self):
            self.offset = 0
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self, size: int) -> bytes:
            if self.offset:
                return b""
            self.offset += 1
            return b"video"

    def fake_urlopen(request):
        nonlocal attempts
        opened_urls.append(request.full_url)
        attempts += 1
        if attempts == 1:
            raise urllib.error.HTTPError(request.full_url, 403, "Forbidden", {}, None)
        return FakeResponse()

    monkeypatch.setattr("urllib.request.urlopen", fake_urlopen)
    download_plan = {
        "dramaId": "drama-1",
        "title": "重试短剧",
        "episodes": [
            {"episodeNo": 1, "sourcePath": "/pan/001.mp4", "size": 5, "downloadUrl": "/files/1.mp4"},
        ],
    }

    files = download_episodes(
        download_plan,
        tmp_path / "drama-1",
        "http://server/api",
        episode_retry_count=2,
        retry_delay_seconds=0,
    )

    episode_file = tmp_path / "drama-1" / "重试短剧-第1集.mp4"
    assert files == [episode_file]
    assert episode_file.read_bytes() == b"video"
    assert not (tmp_path / "drama-1" / "重试短剧-第1集.mp4.part").exists()
    assert opened_urls == ["http://server/files/1.mp4", "http://server/files/1.mp4"]


def test_download_episodes_retries_plain_bad_gateway(tmp_path, monkeypatch):
    attempts = 0

    class FakeResponse:
        headers = {"Content-Length": "5"}

        def __enter__(self):
            self.offset = 0
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self, size: int) -> bytes:
            if self.offset:
                return b""
            self.offset += 1
            return b"video"

    def fake_urlopen(request):
        nonlocal attempts
        attempts += 1
        if attempts == 1:
            raise urllib.error.HTTPError(request.full_url, 502, "Bad Gateway", {}, io.BytesIO(b"bad"))
        return FakeResponse()

    monkeypatch.setattr("urllib.request.urlopen", fake_urlopen)
    download_plan = {
        "dramaId": "drama-1",
        "title": "网关重试",
        "episodes": [
            {"episodeNo": 1, "sourcePath": "/pan/001.mp4", "size": 5, "downloadUrl": "/files/1.mp4"},
        ],
    }

    files = download_episodes(
        download_plan,
        tmp_path / "drama-1",
        "http://server/api",
        episode_retry_count=2,
        retry_delay_seconds=0,
    )

    assert attempts == 2
    assert files == [tmp_path / "drama-1" / "网关重试-第1集.mp4"]


def test_download_episodes_does_not_retry_hongguo_empty_video(tmp_path, monkeypatch):
    attempts = 0

    def fake_urlopen(request):
        nonlocal attempts
        attempts += 1
        body = json.dumps(
            {
                "success": False,
                "data": None,
                "error": {"code": "HONGGUO_VIDEO_EMPTY", "message": "红果播放接口没有返回可下载视频"},
            }
        ).encode("utf-8")
        raise urllib.error.HTTPError(request.full_url, 502, "Bad Gateway", {}, io.BytesIO(body))

    monkeypatch.setattr("urllib.request.urlopen", fake_urlopen)
    download_plan = {
        "dramaId": "drama-1",
        "title": "红果缺集",
        "episodes": [
            {"episodeNo": 25, "sourcePath": "52api://video/25", "size": 0, "downloadUrl": "/files/25.mp4"},
        ],
    }

    with pytest.raises(RuntimeError) as error:
        download_episodes(
            download_plan,
            tmp_path / "drama-1",
            "http://server/api",
            episode_retry_count=3,
            retry_delay_seconds=0,
        )

    assert attempts == 1
    assert "第 25 集下载失败" in str(error.value)
    assert "红果播放接口没有返回可下载视频" in str(error.value)
    assert "HONGGUO_VIDEO_EMPTY" in str(error.value)


def test_download_episodes_skips_allowed_failures(tmp_path, monkeypatch):
    skipped = []

    class FakeResponse:
        headers = {"Content-Length": "5"}

        def __enter__(self):
            self.offset = 0
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self, size: int) -> bytes:
            if self.offset:
                return b""
            self.offset += 1
            return b"video"

    def fake_urlopen(request):
        if request.full_url.endswith("/2.mp4"):
            body = json.dumps(
                {
                    "success": False,
                    "error": {"code": "HONGGUO_VIDEO_EMPTY", "message": "红果播放接口没有返回可下载视频"},
                }
            ).encode("utf-8")
            raise urllib.error.HTTPError(request.full_url, 502, "Bad Gateway", {}, io.BytesIO(body))
        return FakeResponse()

    monkeypatch.setattr("urllib.request.urlopen", fake_urlopen)
    download_plan = {
        "dramaId": "drama-1",
        "title": "允许跳集",
        "episodes": [
            {"episodeNo": 1, "downloadUrl": "/files/1.mp4"},
            {"episodeNo": 2, "downloadUrl": "/files/2.mp4"},
            {"episodeNo": 3, "downloadUrl": "/files/3.mp4"},
            {"episodeNo": 4, "downloadUrl": "/files/4.mp4"},
        ],
    }

    files = download_episodes(
        download_plan,
        tmp_path / "drama-1",
        "http://server/api",
        max_concurrent_downloads=1,
        episode_retry_count=1,
        retry_delay_seconds=0,
        max_skipped_episodes=5,
        skip_callback=lambda index, total, episode, exception: skipped.append((index, str(exception))),
    )

    assert [file.name for file in files] == ["允许跳集-第1集.mp4", "允许跳集-第2集.mp4", "允许跳集-第3集.mp4"]
    assert skipped == [(2, "第 2 集下载失败：红果播放接口没有返回可下载视频（HONGGUO_VIDEO_EMPTY）")]
    metadata = json.loads((tmp_path / "drama-1" / "meta.json").read_text(encoding="utf-8"))
    assert metadata["episodeCount"] == 3
    assert metadata["skippedEpisodeNumbers"] == [2]
    assert [episode["episodeNo"] for episode in metadata["episodes"]] == [1, 2, 3]
    assert metadata["episodes"][1]["originalEpisodeNo"] == 3


def test_download_episodes_fails_when_skipped_failures_exceed_limit(tmp_path, monkeypatch):
    def fake_urlopen(request):
        body = json.dumps(
            {
                "success": False,
                "error": {"code": "HONGGUO_VIDEO_EMPTY", "message": "红果播放接口没有返回可下载视频"},
            }
        ).encode("utf-8")
        raise urllib.error.HTTPError(request.full_url, 502, "Bad Gateway", {}, io.BytesIO(body))

    monkeypatch.setattr("urllib.request.urlopen", fake_urlopen)
    download_plan = {
        "dramaId": "drama-1",
        "title": "超过跳集",
        "episodes": [
            {"episodeNo": episode_no, "downloadUrl": f"/files/{episode_no}.mp4"}
            for episode_no in range(1, 7)
        ],
    }

    with pytest.raises(RuntimeError) as error:
        download_episodes(
            download_plan,
            tmp_path / "drama-1",
            "http://server/api",
            max_concurrent_downloads=1,
            episode_retry_count=1,
            retry_delay_seconds=0,
            max_skipped_episodes=5,
        )

    assert "剧集下载失败超过 5 集" in str(error.value)


def test_download_resources_use_baidu_download_headers(tmp_path, monkeypatch):
    opened_headers = []

    class FakeResponse:
        headers = {"Content-Length": "5"}

        def __enter__(self):
            self.offset = 0
            return self

        def __exit__(self, exc_type, exc, traceback):
            return False

        def read(self, size: int) -> bytes:
            if self.offset:
                return b""
            self.offset += 1
            return b"video"

    def fake_urlopen(request):
        opened_headers.append(dict(request.header_items()))
        return FakeResponse()

    monkeypatch.setattr("urllib.request.urlopen", fake_urlopen)
    download_plan = {
        "dramaId": "drama-1",
        "title": "百度下载头",
        "effectiveCoverUrl": "/covers/drama.jpg",
        "episodes": [
            {"episodeNo": 1, "sourcePath": "/pan/001.mp4", "size": 5, "downloadUrl": "/files/1.mp4"},
        ],
    }

    download_episodes(
        download_plan,
        tmp_path / "drama-1",
        "http://server/api",
        headers={"Authorization": "Bearer token"},
    )

    assert len(opened_headers) == 2
    for headers in opened_headers:
        assert headers["User-agent"] == "pan.baidu.com"
        assert headers["Referer"] == "https://pan.baidu.com/"
        assert headers["Authorization"] == "Bearer token"


def test_download_episodes_downloads_episodes_concurrently_with_limit(tmp_path, monkeypatch):
    active = 0
    max_active = 0
    lock = threading.Lock()
    opened_urls = []

    class FakeResponse:
        def __init__(self, body: bytes, is_video: bool):
            self.body = body
            self.offset = 0
            self.is_video = is_video
            self.entered = False
            self.headers = {"Content-Length": str(len(body))}

        def __enter__(self):
            nonlocal active, max_active
            if self.is_video:
                with lock:
                    active += 1
                    max_active = max(max_active, active)
            self.entered = True
            return self

        def __exit__(self, exc_type, exc, traceback):
            nonlocal active
            if self.is_video and self.entered:
                with lock:
                    active -= 1
            return False

        def read(self, size: int) -> bytes:
            time.sleep(0.02)
            if self.offset >= len(self.body):
                return b""
            chunk = self.body[self.offset : self.offset + size]
            self.offset += len(chunk)
            return chunk

    def fake_urlopen(request):
        opened_urls.append(request.full_url)
        is_video = "/files/" in request.full_url
        return FakeResponse(request.full_url.encode(), is_video)

    monkeypatch.setattr("urllib.request.urlopen", fake_urlopen)
    download_plan = {
        "dramaId": "drama-1",
        "title": "并发短剧",
        "episodes": [
            {"episodeNo": episode_no, "downloadUrl": f"/files/{episode_no}.mp4"}
            for episode_no in range(1, 9)
        ],
    }

    files = download_episodes(download_plan, tmp_path / "drama-1", "http://server/api")

    assert files == [tmp_path / "drama-1" / f"并发短剧-第{episode_no}集.mp4" for episode_no in range(1, 9)]
    assert max_active == 6
    assert {url for url in opened_urls if "/files/" in url} == {
        f"http://server/files/{episode_no}.mp4" for episode_no in range(1, 9)
    }


def test_episode_video_filename_uses_ai_title_and_episode_number():
    assert (
        episode_video_filename(
            {"title": "原始剧名", "aiTitle": "AI剧名"},
            {"episodeNo": 12},
            1,
        )
        == "AI剧名-第12集.mp4"
    )
